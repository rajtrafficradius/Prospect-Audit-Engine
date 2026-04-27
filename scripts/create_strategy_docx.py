#!/usr/bin/env python3
"""
Prospect Audit & Strategy — Branded DOCX Generator (SEO + AEO + GEO)
=====================================================================
Generates a 15-25 page branded strategy document with 12 sections
covering all three layers of modern search.

Adapt this template for each prospect by replacing placeholder data
with real data from the JSON/CSV files generated in Phases 1-5.

Usage: python3 create_strategy_docx.py
Requires: python-docx (pip install python-docx)
"""

import os
import json
import re
import math
import textwrap
import hashlib
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from PIL import Image, ImageDraw, ImageFont

# ══════════════════════════════════════════════════════════════
# CONFIGURATION — Update these for each prospect
# ══════════════════════════════════════════════════════════════

import datetime
SCRIPT_DIR = os.path.abspath(os.path.dirname(__file__))
DEFAULT_OUTPUT_ROOT = os.path.abspath(os.path.join(SCRIPT_DIR, "..", "output"))
OUTPUT_ROOT = os.path.abspath(os.environ.get("OUTPUT_DIR", DEFAULT_OUTPUT_ROOT))
PROSPECT_NAME = os.environ.get("PROSPECT_NAME", "TrafficRadius Prospect")
PROSPECT_DOMAIN = os.environ.get("PROSPECT_DOMAIN", "example.com")
PROSPECT_DATE = os.environ.get("PROSPECT_DATE", datetime.datetime.now().strftime("%B %d, %Y"))
OUTPUT_DIR = os.path.join(OUTPUT_ROOT, "deliverables")
CHARTS_DIR = os.path.join(OUTPUT_ROOT, "charts")
LOGO_PATH = os.path.join(os.path.dirname(__file__), '..', 'src', 'static', 'logo.png')
DATA_DIR = OUTPUT_ROOT

try:
    with open(os.path.join(DATA_DIR, "strategy_narrative.json"), "r") as f:
        narrative_data = json.load(f)
except Exception:
    narrative_data = {}

try:
    with open(os.path.join(DATA_DIR, "market_intelligence.json"), "r") as f:
        market_data = json.load(f)
except Exception:
    market_data = {}

try:
    with open(os.path.join(DATA_DIR, "audit_findings.json"), "r") as f:
        audit_data = json.load(f)
except Exception:
    audit_data = {}

try:
    with open(os.path.join(DATA_DIR, "cro_assessment.json"), "r") as f:
        cro_data = json.load(f)
except Exception:
    cro_data = {}

try:
    with open(os.path.join(DATA_DIR, "business_analysis.json"), "r") as f:
        business_data = json.load(f)
except Exception:
    business_data = {}

try:
    with open(os.path.join(DATA_DIR, "competitor_shadowing.json"), "r") as f:
        shadow_data = json.load(f)
except Exception:
    shadow_data = {}

# ── Brand Colors ──────────────────────────────────────────────
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
BLUE = RGBColor(0x2E, 0x50, 0x90)
LIGHT_BLUE = RGBColor(0x4A, 0x90, 0xD9)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GREY = RGBColor(0x2D, 0x37, 0x48)
MID_GREY = RGBColor(0x4A, 0x55, 0x68)
LIGHT_GREY = RGBColor(0x71, 0x80, 0x96)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Layer accent colors
SEO_BLUE = RGBColor(0x2E, 0x50, 0x90)
AEO_ORANGE = RGBColor(0xE6, 0x7E, 0x22)
GEO_PURPLE = RGBColor(0x9B, 0x59, 0xB6)

NAVY_HEX = "1B2A4A"
BLUE_HEX = "2E5090"
GREEN_HEX = "2E7D32"
BLACK_HEX = "000000"
AEO_HEX = "E67E22"
GEO_HEX = "9B59B6"
LIGHT_BG = "F7FAFC"
TABLE_HEADER_BG = "1B2A4A"
TABLE_ALT_ROW = "F0F2F5"
DOCX_IMAGE_DPI = (300, 300)
DOCX_CHART_DPI = 450
TRAFFICRADIUS_EMAIL = "Info@trafficradius.com.au"
TRAFFICRADIUS_PHONE = "+61 1300 852 340"
TRAFFICRADIUS_WEB = "trafficradius.com.au"

# ══════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ══════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "auto")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def apply_full_grid_borders(table, color=BLACK_HEX, sz="10"):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                left={"sz": sz, "color": color, "val": "single"},
                right={"sz": sz, "color": color, "val": "single"},
                top={"sz": sz, "color": color, "val": "single"},
                bottom={"sz": sz, "color": color, "val": "single"},
            )

def add_branded_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, TABLE_HEADER_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_GREY
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_ROW)
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    apply_full_grid_borders(table, color=BLACK_HEX, sz="10")
    doc.add_paragraph("")
    return table

def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.style = doc.styles['Heading 1']
        p.clear()
        run = p.add_run(text)
        run.font.color.rgb = NAVY
        run.font.size = Pt(24)
        run.font.name = 'Calibri'
        run.bold = True
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
    elif level == 2:
        p.style = doc.styles['Heading 2']
        p.clear()
        run = p.add_run(text)
        run.font.color.rgb = BLUE
        run.font.size = Pt(18)
        run.font.name = 'Calibri'
        run.bold = True
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 3:
        p.style = doc.styles['Heading 3']
        p.clear()
        run = p.add_run(text)
        run.font.color.rgb = DARK_GREY
        run.font.size = Pt(14)
        run.font.name = 'Calibri'
        run.bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    return p

def add_body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color if color else DARK_GREY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_callout(doc, text, accent_color_hex=BLUE_HEX):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK_GREY
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_cell_shading(cell, LIGHT_BG)
    set_cell_border(cell,
        left={"sz": "24", "color": accent_color_hex, "val": "single"},
        top={"sz": "4", "color": "E2E8F0", "val": "single"},
        bottom={"sz": "4", "color": "E2E8F0", "val": "single"},
        right={"sz": "4", "color": "E2E8F0", "val": "single"}
    )
    doc.add_paragraph("")

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="{BLUE_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

def _caption_from_filename(filename):
    base = os.path.splitext(os.path.basename(filename))[0]
    text = base.replace("_", " ").replace("-", " ").strip().title()
    return f"Figure: {text}"

def optimize_image_for_display(path, width_inches, dpi=300):
    if not path or not os.path.exists(path):
        return path
    try:
        os.makedirs(os.path.join(DATA_DIR, ".docx_optimized"), exist_ok=True)
        with Image.open(path) as img:
            if img.mode in ("RGBA", "LA") or (img.mode == "P" and "transparency" in img.info):
                alpha_img = img.convert("RGBA")
                white_bg = Image.new("RGBA", alpha_img.size, (255, 255, 255, 255))
                white_bg.alpha_composite(alpha_img)
                img = white_bg.convert("RGB")
            else:
                img = img.convert("RGB")
            target_w = max(1, int(round(width_inches * dpi)))
            aspect = img.height / max(img.width, 1)
            target_h = max(1, int(round(target_w * aspect)))
            if img.width != target_w or img.height != target_h:
                img = img.resize((target_w, target_h), Image.Resampling.LANCZOS)
            base = os.path.splitext(os.path.basename(path))[0]
            out_path = os.path.join(DATA_DIR, ".docx_optimized", f"{base}_{target_w}w.png")
            img.save(out_path, dpi=(dpi, dpi))
            return out_path
    except Exception:
        return path

def add_chart_image(doc, filename, width=6.0, caption=None):
    path = os.path.join(CHARTS_DIR, filename)
    if os.path.exists(path):
        optimized_path = optimize_image_for_display(path, width)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        run = p.add_run()
        run.add_picture(optimized_path, width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.keep_together = True
        cap.paragraph_format.space_before = Pt(2)
        cap.paragraph_format.space_after = Pt(8)
        cap_text = caption if caption else _caption_from_filename(filename)
        cap_run = cap.add_run(cap_text)
        cap_run.font.name = "Calibri"
        cap_run.font.size = Pt(9.5)
        cap_run.font.color.rgb = LIGHT_GREY
        return p
    return None

def add_chart_with_caption(doc, filename, caption, width=6.0):
    return bool(add_chart_image(doc, filename, width=width, caption=caption))

def add_centered_image(doc, path, width=2.8):
    if not path or not os.path.exists(path):
        return None
    optimized_path = optimize_image_for_display(path, width)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(optimized_path, width=Inches(width))
    p.paragraph_format.space_after = Pt(8)
    return p

def add_bullet_list(doc, items, color=None):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(str(item))
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.font.color.rgb = color if color else DARK_GREY
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("")

def create_action_plan_roadmap_image(rows):
    out_path = os.path.join(CHARTS_DIR, "chapter11_action_plan_roadmap.png")
    os.makedirs(CHARTS_DIR, exist_ok=True)
    canvas_w, canvas_h = 1800, 900
    img = Image.new("RGB", (canvas_w, canvas_h), (249, 251, 255))
    draw = ImageDraw.Draw(img)

    title_font = _load_dashboard_font(56, bold=True)
    sub_font = _load_dashboard_font(28, bold=False)
    phase_font = _load_dashboard_font(30, bold=True)
    label_font = _load_dashboard_font(24, bold=True)
    body_font = _load_dashboard_font(19, bold=False)
    small_font = _load_dashboard_font(18, bold=False)

    border = (225, 232, 243)
    white_rgb = (255, 255, 255)
    navy_rgb = (27, 42, 74)
    blue_rgb = (36, 99, 235)
    sky_rgb = (186, 208, 251)
    text_mid = (86, 96, 112)
    phase_accents = [(18, 58, 138), (36, 99, 235), (98, 126, 234)]

    draw.rounded_rectangle((25, 25, canvas_w - 25, canvas_h - 25), radius=34, fill=white_rgb, outline=border, width=3)
    draw.text((80, 80), "Action Plan Roadmap", font=title_font, fill=navy_rgb)
    draw.text((82, 150), "A phased execution flow across SEO, AEO, and GEO priorities.", font=sub_font, fill=text_mid)
    draw.rounded_rectangle((80, 205, 250, 214), radius=4, fill=blue_rgb)

    x1, x2, y = 160, canvas_w - 170, 315
    draw.line((x1, y, x2, y), fill=blue_rgb, width=8)
    draw.polygon([(x2, y), (x2 - 28, y - 18), (x2 - 28, y + 18)], fill=blue_rgb)

    centers = [360, 900, 1440]
    day_labels = ["Days 1-30", "Days 31-60", "Days 61-90"]

    for idx, row in enumerate(rows[:3]):
        cx = centers[idx]
        draw.ellipse((cx - 30, y - 30, cx + 30, y + 30), fill=white_rgb, outline=sky_rgb, width=4)
        draw.ellipse((cx - 18, y - 18, cx + 18, y + 18), fill=blue_rgb)
        draw.text((cx - 44, 250), day_labels[idx], font=small_font, fill=text_mid)
        draw.text((cx - 58, 370), f"Phase {idx + 1}", font=phase_font, fill=blue_rgb)
        draw.text((cx - 130, 408), row["phase"], font=label_font, fill=navy_rgb)

        card_w, card_h = 470, 315
        left = int(cx - card_w / 2)
        top = 470
        draw.rounded_rectangle((left + 10, top + 12, left + card_w + 10, top + card_h + 12), radius=28, fill=(245, 247, 250))
        draw.rounded_rectangle((left, top, left + card_w, top + card_h), radius=28, fill=white_rgb, outline=(218, 225, 234), width=2)
        accent = phase_accents[idx % len(phase_accents)]
        draw.rounded_rectangle((left, top, left + card_w, top + 16), radius=28, fill=accent)
        draw.rectangle((left, top + 16, left + card_w, top + 38), fill=accent)

        section_y = top + 42
        for j, (label, value) in enumerate([("SEO", row["seo"]), ("AEO", row["aeo"]), ("GEO", row["geo"])]):
            draw.text((left + 24, section_y), label, font=label_font, fill=navy_rgb)
            wrapped = textwrap.fill(str(value), width=42).splitlines()[:2]
            draw.multiline_text((left + 24, section_y + 32), "\n".join(wrapped), font=body_font, fill=text_mid, spacing=4)
            if j < 2:
                draw.line((left + 22, section_y + 82, left + card_w - 22, section_y + 82), fill=border, width=2)
            section_y += 88

    footer = "Use this roadmap alongside the detailed action table to align owners, milestones, and expected business outcomes."
    draw.rounded_rectangle((260, 812, canvas_w - 260, 866), radius=18, fill=(244, 247, 253), outline=sky_rgb, width=2)
    footer_bbox = draw.textbbox((0, 0), footer, font=small_font)
    footer_x = int((canvas_w - (footer_bbox[2] - footer_bbox[0])) / 2)
    draw.text((footer_x, 829), footer, font=small_font, fill=text_mid)

    img.save(out_path, dpi=DOCX_IMAGE_DPI)
    return out_path

def add_action_plan_roadmap(doc, phase_rows):
    add_heading_styled(doc, "Execution Guidance", 2)
    add_bullet_list(doc, [
        "Phase 1 focuses on removing technical blockers quickly so search engines and answer engines can crawl, interpret, and trust core service pages from day one.",
        "Phase 2 expands answer-ready content and entity signals so the site can convert visibility into qualified enquiries across SEO, AEO, and GEO surfaces.",
        "Phase 3 scales authority, measurement, and outreach so revenue gains are sustained rather than dependent on one-off optimizations.",
    ], color=MID_GREY)

    add_heading_styled(doc, "Roadmap Summary", 2)
    roadmap_path = create_action_plan_roadmap_image(phase_rows)
    add_chart_image(doc, os.path.basename(roadmap_path), width=6.9, caption="The roadmap below translates the action plan into sequenced milestones across the implementation window, making the delivery flow easier to review at a glance.")

    add_heading_styled(doc, "Success Measures", 2)
    add_bullet_list(doc, [
        "Operational milestone: technical fixes, schema deployment, and AI-crawl readiness should be completed early so later phases can build on a stable foundation.",
        "Commercial milestone: new answer-engine content and service-page upgrades should begin producing stronger enquiry quality, not just higher impressions.",
        "Leadership milestone: monthly reporting should connect visibility gains to leads, pipeline value, and implementation velocity so decision-making stays commercial.",
    ], color=MID_GREY)

def add_trafficradius_next_steps_extension(doc):
    add_heading_styled(doc, "TrafficRadius Partnership Snapshot", 2)
    add_centered_image(doc, LOGO_PATH, width=2.8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_parts = [
        ("Email: ", True, DARK_GREY),
        (TRAFFICRADIUS_EMAIL, False, MID_GREY),
        ("   |   Phone: ", True, DARK_GREY),
        (TRAFFICRADIUS_PHONE, False, MID_GREY),
        ("   |   Web: ", True, DARK_GREY),
        (TRAFFICRADIUS_WEB, False, MID_GREY),
    ]
    for txt, is_bold, color in contact_parts:
        run = p.add_run(txt)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.bold = is_bold
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    run = p.add_run("How this partnership translates into action")
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = DARK_GREY
    p.paragraph_format.space_after = Pt(4)

    add_bullet_list(doc, [
        "TrafficRadius brings one integrated operating model across SEO, AEO, CRO, and GEO so execution stays aligned with revenue goals instead of fragmented across channels.",
        "A senior strategy team can use the roadmap, action plan, and monthly KPI reviews to prioritise the highest-impact fixes first and accelerate visible wins early.",
        "The first working sessions can focus on implementation sequencing, reporting cadence, ownership, and success metrics so the engagement starts with clarity and momentum.",
    ], color=MID_GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ready to get started? Let's build this together.")
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A simple weekly KPI tracker can support accountability across technical completion, content production, authority wins, and lead growth during the first 90 days.")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = LIGHT_GREY
    p.paragraph_format.space_after = Pt(4)

    add_heading_styled(doc, "Engagement Model", 2)
    add_body(doc,
        "A practical first-month engagement model helps translate strategy into visible execution without losing momentum between planning and delivery.",
        color=MID_GREY
    )
    add_bullet_list(doc, [
        "Week 1 example: kickoff workshop, KPI alignment, access collection, and prioritisation of urgent technical blockers.",
        "Week 2 example: service-page upgrades, answer-engine content mapping, and rollout sequencing for high-value schema opportunities.",
        "Week 3-4 example: authority-building outreach, reporting dashboard setup, and leadership review of early wins plus next-sprint priorities.",
    ], color=MID_GREY)

def add_brand_dashboard_panel(doc, title, cards):
    """
    Add a compact visual dashboard panel using brand colors (Blue/Black/Green).
    cards: list of (label, value)
    """
    add_heading_styled(doc, title, 2)
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    try:
        for row in table.rows:
            row.cells[0].width = Inches(3.2)
            row.cells[1].width = Inches(3.2)
    except Exception:
        pass
    palette = [BLUE_HEX, BLACK_HEX, GREEN_HEX, BLUE_HEX]
    for i in range(2):
        for j in range(2):
            idx = i * 2 + j
            label, value = cards[idx] if idx < len(cards) else ("Metric", "N/A")
            cell = table.rows[i].cells[j]
            cell.text = ""
            set_cell_shading(cell, palette[idx])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p.add_run(f"{label}\n")
            r1.font.name = "Calibri"
            r1.font.size = Pt(9.5)
            r1.font.bold = True
            r1.font.color.rgb = WHITE
            r2 = p.add_run(str(value))
            r2.font.name = "Calibri"
            r2.font.size = Pt(12)
            r2.font.bold = True
            r2.font.color.rgb = WHITE
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    apply_full_grid_borders(table, color=BLACK_HEX, sz="10")
    doc.add_paragraph("")

def _load_dashboard_font(size, bold=False):
    candidates = [
        "arialbd.ttf" if bold else "arial.ttf",
        "calibrib.ttf" if bold else "calibri.ttf",
        "segoeuib.ttf" if bold else "segoeui.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except Exception:
            continue
    return ImageFont.load_default()

def _draw_target_icon(draw, box, color):
    x0, y0, x1, y1 = box
    cx = (x0 + x1) / 2
    cy = (y0 + y1) / 2
    r1 = min(x1 - x0, y1 - y0) * 0.34
    r2 = r1 * 0.68
    r3 = r1 * 0.34
    draw.ellipse((cx-r1, cy-r1, cx+r1, cy+r1), outline=color, width=3)
    draw.ellipse((cx-r2, cy-r2, cx+r2, cy+r2), outline=color, width=3)
    draw.ellipse((cx-r3, cy-r3, cx+r3, cy+r3), outline=color, width=3)
    draw.line((cx, cy, cx + r1 * 0.9, cy - r1 * 0.9), fill=color, width=3)
    draw.polygon([
        (cx + r1 * 0.9, cy - r1 * 0.9),
        (cx + r1 * 1.18, cy - r1 * 1.02),
        (cx + r1 * 1.02, cy - r1 * 0.72),
    ], outline=color, fill=None, width=3)

def _draw_chart_icon(draw, box, color):
    x0, y0, x1, y1 = box
    left = x0 + 12
    bottom = y1 - 16
    bar_w = (x1 - x0) * 0.1
    heights = [22, 36, 28, 48, 60]
    for idx, h in enumerate(heights):
        bx = left + idx * (bar_w + 12)
        draw.rectangle((bx, bottom - h, bx + bar_w, bottom), outline=color, width=3)
    points = [
        (left - 8, bottom - 8),
        (left + 20, bottom - 34),
        (left + 56, bottom - 18),
        (left + 96, bottom - 44),
        (left + 136, bottom - 72),
    ]
    draw.line(points, fill=color, width=3)
    draw.polygon([
        (points[-1][0], points[-1][1]),
        (points[-1][0] - 18, points[-1][1] + 4),
        (points[-1][0] - 8, points[-1][1] + 18),
    ], outline=color, fill=None, width=3)

def _draw_coin_icon(draw, box, color):
    x0, y0, x1, y1 = box
    w = x1 - x0
    h = y1 - y0
    for offset in [22, 12, 2]:
        draw.ellipse((x0 + w*0.28, y0 + offset, x0 + w*0.72, y0 + offset + h*0.16), outline=color, width=3)
        draw.rectangle((x0 + w*0.28, y0 + offset + h*0.08, x0 + w*0.72, y0 + offset + h*0.28), outline=color, width=3)
        draw.arc((x0 + w*0.28, y0 + offset + h*0.20, x0 + w*0.72, y0 + offset + h*0.36), start=0, end=180, fill=color, width=3)
    cx = x0 + w * 0.50
    cy = y0 + h * 0.36
    r = min(w, h) * 0.13
    draw.ellipse((cx-r, cy-r, cx+r, cy+r), outline=color, width=3)
    draw.text((cx-8, cy-12), "$", font=_load_dashboard_font(20, bold=True), fill=color)

def _draw_gauge_icon(draw, box, color):
    x0, y0, x1, y1 = box
    cx = (x0 + x1) / 2
    cy = y1 - 12
    r = min(x1 - x0, y1 - y0) * 0.44
    draw.arc((cx-r, cy-r, cx+r, cy+r), start=180, end=360, fill=color, width=3)
    for angle in [190, 220, 250, 280, 310, 340]:
        radians = math.radians(angle)
        x_start = cx + (r - 12) * math.cos(radians)
        y_start = cy + (r - 12) * math.sin(radians)
        x_end = cx + r * math.cos(radians)
        y_end = cy + r * math.sin(radians)
        draw.line((x_start, y_start, x_end, y_end), fill=color, width=3)
    needle_angle = math.radians(325)
    draw.line((cx, cy, cx + (r - 18) * math.cos(needle_angle), cy + (r - 18) * math.sin(needle_angle)), fill=color, width=4)
    draw.ellipse((cx-8, cy-8, cx+8, cy+8), outline=color, width=3)

def create_executive_snapshot_dashboard(cards):
    out_path = os.path.join(CHARTS_DIR, "executive_snapshot_dashboard.png")
    os.makedirs(CHARTS_DIR, exist_ok=True)
    canvas_w, canvas_h = 2360, 760
    img = Image.new("RGB", (canvas_w, canvas_h), (255, 255, 255))
    draw = ImageDraw.Draw(img)

    section_font = _load_dashboard_font(44, bold=True)
    title_font = _load_dashboard_font(34, bold=True)
    label_font = _load_dashboard_font(24, bold=False)
    value_font = _load_dashboard_font(54, bold=True)
    note_font = _load_dashboard_font(20, bold=False)
    border = (214, 221, 230)
    shadow = (244, 247, 251)
    label_color = (96, 104, 118)
    value_color = (49, 61, 79)
    divider = (227, 234, 242)

    card_w = 548
    card_h = 330
    gap_x = 34
    top = 250
    left = 40
    header_left = 46
    accent_colors = [
        (46, 80, 144),
        (56, 159, 195),
        (232, 127, 33),
        (142, 91, 183),
    ]
    notes = [
        "Organic search footprint",
        "Current monthly discovery baseline",
        "Estimated monetized search demand",
        "Integrated readiness benchmark",
    ]

    draw.text((header_left, 28), "Executive Snapshot Dashboard", font=section_font, fill=(46, 80, 144))
    #draw.text((header_left, 128), "Executive Snapshot Dashboard", font=title_font, fill=(36, 50, 76))
    draw.line((header_left, 210, canvas_w - 44, 210), fill=divider, width=3)

    for idx in range(4):
        x = left + idx * (card_w + gap_x)
        y = top
        draw.rounded_rectangle((x + 7, y + 8, x + card_w + 7, y + card_h + 8), radius=24, fill=shadow)
        draw.rounded_rectangle((x, y, x + card_w, y + card_h), radius=24, fill=(255, 255, 255), outline=border, width=3)
        accent = accent_colors[idx % len(accent_colors)]
        draw.rounded_rectangle((x, y, x + card_w, y + 18), radius=24, fill=accent)
        draw.rectangle((x, y + 18, x + card_w, y + 40), fill=accent)

        label, value = cards[idx] if idx < len(cards) else ("Metric", "N/A")
        value_str = str(value)
        if "value" in str(label).lower():
            value_str = _format_money_display(value)
        elif str(value).replace(",", "").isdigit():
            value_str = _format_int_display(value)

        draw.text((x + 34, y + 78), str(label), font=label_font, fill=label_color)
        draw.text((x + 34, y + 144), value_str, font=value_font, fill=value_color)
        note_lines = textwrap.fill(notes[idx], width=22).splitlines()[:2]
        draw.multiline_text((x + 34, y + 236), "\n".join(note_lines), font=note_font, fill=(116, 126, 140), spacing=4)

    img.save(out_path, dpi=DOCX_IMAGE_DPI)
    return out_path

def add_executive_snapshot_dashboard(doc, title, cards):
    heading = add_heading_styled(doc, title, 2)
    heading.paragraph_format.keep_with_next = True
    image_path = create_executive_snapshot_dashboard(cards)
    optimized_path = optimize_image_for_display(image_path, 6.6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    run = p.add_run()
    run.add_picture(optimized_path, width=Inches(6.6))
    doc.add_paragraph("")

def create_integrated_layer_visual_dashboard(cards):
    out_path = os.path.join(CHARTS_DIR, "integrated_layer_cards.png")
    os.makedirs(CHARTS_DIR, exist_ok=True)
    canvas_w, canvas_h = 2360, 760
    img = Image.new("RGB", (canvas_w, canvas_h), (255, 255, 255))
    draw = ImageDraw.Draw(img)

    section_font = _load_dashboard_font(44, bold=True)
    title_font = _load_dashboard_font(34, bold=True)
    label_font = _load_dashboard_font(24, bold=False)
    value_font = _load_dashboard_font(54, bold=True)
    note_font = _load_dashboard_font(20, bold=False)
    border = (214, 221, 230)
    shadow = (244, 247, 251)
    label_color = (96, 104, 118)
    value_color = (49, 61, 79)
    divider = (227, 234, 242)

    card_w = 548
    card_h = 330
    gap_x = 34
    top = 250
    left = 40
    header_left = 46
    accent_colors = [
        (46, 80, 144),
        (56, 159, 195),
        (232, 127, 33),
        (142, 91, 183),
    ]

    draw.text((header_left, 28), "Integrated Layer Visual Dashboard", font=section_font, fill=(46, 80, 144))
    #draw.text((header_left, 128), "Integrated Layer Visual Dashboard", font=title_font, fill=(36, 50, 76))
    draw.line((header_left, 210, canvas_w - 44, 210), fill=divider, width=3)

    for idx in range(4):
        x = left + idx * (card_w + gap_x)
        y = top
        draw.rounded_rectangle((x + 7, y + 8, x + card_w + 7, y + card_h + 8), radius=24, fill=shadow)
        draw.rounded_rectangle((x, y, x + card_w, y + card_h), radius=24, fill=(255, 255, 255), outline=border, width=3)
        accent = accent_colors[idx % len(accent_colors)]
        draw.rounded_rectangle((x, y, x + card_w, y + 18), radius=24, fill=accent)
        draw.rectangle((x, y + 18, x + card_w, y + 40), fill=accent)

        card = cards[idx] if idx < len(cards) else ("Metric", "N/A", "", accent)
        label = card[0]
        value = card[1]
        note = card[2] if len(card) > 2 else ""
        value_str = _format_int_display(value) if str(value).replace(",", "").isdigit() else str(value)

        draw.text((x + 34, y + 78), str(label), font=label_font, fill=label_color)
        draw.text((x + 34, y + 144), value_str, font=value_font, fill=value_color)
        note_lines = textwrap.fill(str(note), width=22).splitlines()[:2]
        draw.multiline_text((x + 34, y + 236), "\n".join(note_lines), font=note_font, fill=(116, 126, 140), spacing=4)

    img.save(out_path, dpi=DOCX_IMAGE_DPI)
    return out_path

def add_integrated_layer_visual_dashboard(doc):
    cards = build_integrated_layer_cards()
    image_path = create_integrated_layer_visual_dashboard(cards)
    optimized_path = optimize_image_for_display(image_path, 6.6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    run = p.add_run()
    run.add_picture(optimized_path, width=Inches(6.6))
    doc.add_paragraph("")

def create_opportunity_mix_dashboard_image(cards):
    out_path = os.path.join(CHARTS_DIR, "opportunity_mix_cards.png")
    os.makedirs(CHARTS_DIR, exist_ok=True)
    canvas_w, canvas_h = 2360, 760
    img = Image.new("RGB", (canvas_w, canvas_h), (255, 255, 255))
    draw = ImageDraw.Draw(img)

    section_font = _load_dashboard_font(44, bold=True)
    title_font = _load_dashboard_font(34, bold=True)
    label_font = _load_dashboard_font(22, bold=False)
    value_font = _load_dashboard_font(46, bold=True)
    note_font = _load_dashboard_font(18, bold=False)
    border = (214, 221, 230)
    shadow = (244, 247, 251)
    label_color = (96, 104, 118)
    value_color = (49, 61, 79)
    divider = (227, 234, 242)

    card_w = 548
    card_h = 330
    gap_x = 34
    top = 250
    left = 40
    header_left = 46
    accent_colors = [
        (46, 80, 144),
        (232, 127, 33),
        (142, 91, 183),
        (46, 125, 50),
    ]

    draw.text((header_left, 28), "Opportunity Mix Dashboard", font=section_font, fill=(46, 80, 144))
    #draw.text((header_left, 128), "Opportunity Mix Dashboard", font=title_font, fill=(36, 50, 76))
    draw.text((header_left, 176), "Executive KPI view of the highest-value opportunity mix across modern search layers", font=note_font, fill=(116, 126, 140))
    draw.line((header_left, 210, canvas_w - 44, 210), fill=divider, width=3)

    for idx in range(4):
        x = left + idx * (card_w + gap_x)
        y = top
        draw.rounded_rectangle((x + 7, y + 8, x + card_w + 7, y + card_h + 8), radius=24, fill=shadow)
        draw.rounded_rectangle((x, y, x + card_w, y + card_h), radius=24, fill=(255, 255, 255), outline=border, width=3)
        accent = accent_colors[idx % len(accent_colors)]
        draw.rounded_rectangle((x, y, x + card_w, y + 18), radius=24, fill=accent)
        draw.rectangle((x, y + 18, x + card_w, y + 40), fill=accent)

        card = cards[idx] if idx < len(cards) else ("Metric", "N/A", "", accent)
        label = card[0]
        value = card[1]
        note = card[2] if len(card) > 2 else ""
        value_str = str(value)
        draw.text((x + 34, y + 76), str(label), font=label_font, fill=label_color)
        draw.text((x + 34, y + 136), value_str, font=value_font, fill=value_color)
        note_lines = textwrap.fill(str(note), width=21).splitlines()[:3]
        draw.multiline_text((x + 34, y + 228), "\n".join(note_lines), font=note_font, fill=(116, 126, 140), spacing=4)

    img.save(out_path, dpi=DOCX_IMAGE_DPI)
    return out_path

def add_opportunity_mix_dashboard(doc):
    cards = build_opportunity_mix_cards()
    image_path = create_opportunity_mix_dashboard_image(cards)
    optimized_path = optimize_image_for_display(image_path, 6.6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    run = p.add_run()
    run.add_picture(optimized_path, width=Inches(6.6))
    doc.add_paragraph("")

def _normalize_card_spec(card):
    if isinstance(card, dict):
        return {
            "label": str(card.get("label", "Metric")),
            "value": str(card.get("value", "N/A")),
            "note": str(card.get("note", "")),
            "accent": tuple(card.get("accent", (46, 80, 144))),
        }
    if len(card) >= 4:
        label, value, note, accent = card[:4]
        return {"label": str(label), "value": str(value), "note": str(note), "accent": tuple(accent)}
    if len(card) == 3:
        label, value, note = card
        return {"label": str(label), "value": str(value), "note": str(note), "accent": (46, 80, 144)}
    label, value = card[:2]
    return {"label": str(label), "value": str(value), "note": "", "accent": (46, 80, 144)}

def create_modern_card_dashboard(title, cards, filename, subtitle="", columns=4):
    out_path = os.path.join(CHARTS_DIR, filename)
    os.makedirs(CHARTS_DIR, exist_ok=True)
    card_specs = [_normalize_card_spec(card) for card in cards]
    columns = max(1, min(columns, len(card_specs) if card_specs else 1))
    rows = max(1, math.ceil(len(card_specs) / columns))

    canvas_w = 3200
    header_h = 250 if subtitle else 190
    side_margin = 72
    gap_x = 34
    gap_y = 34
    card_w = int((canvas_w - side_margin * 2 - gap_x * (columns - 1)) / columns)
    card_h = 320
    canvas_h = header_h + rows * card_h + (rows - 1) * gap_y + 90

    img = Image.new("RGB", (canvas_w, canvas_h), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    title_font = _load_dashboard_font(42, bold=True)
    subtitle_font = _load_dashboard_font(22, bold=False)
    label_font = _load_dashboard_font(23, bold=True)
    value_font = _load_dashboard_font(40, bold=True)
    note_font = _load_dashboard_font(18, bold=False)

    draw.text((side_margin, 48), title, font=title_font, fill=(27, 42, 74))
    if subtitle:
        draw.text((side_margin, 114), subtitle, font=subtitle_font, fill=(98, 111, 132))
    draw.line((side_margin, header_h - 26, canvas_w - side_margin, header_h - 26), fill=(225, 232, 241), width=3)

    for idx, spec in enumerate(card_specs):
        row = idx // columns
        col = idx % columns
        x = side_margin + col * (card_w + gap_x)
        y = header_h + row * (card_h + gap_y)
        accent = tuple(spec["accent"])
        draw.rounded_rectangle((x + 10, y + 12, x + card_w + 10, y + card_h + 12), radius=28, fill=(245, 247, 250))
        draw.rounded_rectangle((x, y, x + card_w, y + card_h), radius=28, fill=(255, 255, 255), outline=(218, 225, 234), width=2)
        draw.rounded_rectangle((x, y, x + card_w, y + 16), radius=28, fill=accent)
        draw.rectangle((x, y + 16, x + card_w, y + 38), fill=accent)
        draw.text((x + 28, y + 58), spec["label"], font=label_font, fill=(86, 98, 116))
        draw.text((x + 28, y + 110), spec["value"], font=value_font, fill=(26, 34, 50))
        if spec["note"]:
            wrapped = textwrap.fill(spec["note"], width=max(20, int(card_w / 26)))
            draw.multiline_text((x + 28, y + 185), wrapped, font=note_font, fill=(96, 108, 126), spacing=6)

    img.save(out_path, dpi=DOCX_IMAGE_DPI)
    return out_path

def add_modern_card_dashboard(doc, title, cards, filename, subtitle="", columns=4, width=6.6):
    image_path = create_modern_card_dashboard(title, cards, filename, subtitle=subtitle, columns=columns)
    optimized_path = optimize_image_for_display(image_path, width)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(optimized_path, width=Inches(width))
    doc.add_paragraph("")

def build_integrated_layer_cards():
    scorecard = audit_data.get("scorecard", {})
    cards = [
        ("SEO Score", scorecard.get("seo_score", "N/A"), "Traditional search readiness", (46, 80, 144)),
        ("AEO Score", scorecard.get("aeo_score", "N/A"), "Answer engine readiness", (230, 126, 34)),
        ("GEO Score", scorecard.get("geo_score", "N/A"), "Generative engine visibility", (155, 89, 182)),
        ("CRO Score", scorecard.get("cro_score", "N/A"), "Conversion readiness baseline", (46, 125, 50)),
    ]
    return cards

def build_opportunity_mix_cards():
    seo_label = "Transactional SEO" if market_data.get("prospect", {}).get("top_keywords") else "SEO Opportunity"
    aeo_label = "Question Intent" if market_data.get("aeo_indicators", {}).get("question_keywords_found", 0) else "AEO Opportunity"
    geo_label = "Informational Queries" if market_data.get("geo_indicators", {}).get("informational_keywords_found", 0) else "GEO Opportunity"
    seo_value = int(market_data.get("prospect", {}).get("organic_traffic_value", 0) or 0)
    aeo_value = sum(int(k.get("Search Volume", 0) or 0) for k in market_data.get("aeo_indicators", {}).get("top_question_keywords", []))
    geo_value = sum(int(k.get("Search Volume", 0) or 0) for k in market_data.get("geo_indicators", {}).get("top_informational_keywords", []))
    mix = {
        seo_label: seo_value,
        aeo_label: aeo_value,
        geo_label: geo_value,
    }
    primary_focus = max(mix, key=mix.get) if mix else "High Value Opportunities"
    return [
        ("Top SEO Cluster", seo_label, "Most established search-demand layer", (46, 80, 144)),
        ("Top AEO Cluster", aeo_label, "Question-led answer opportunity", (230, 126, 34)),
        ("Top GEO Cluster", geo_label, "AI discovery-led query cluster", (155, 89, 182)),
        ("Primary Focus", primary_focus, "Highest commercial upside based on current mix", (46, 125, 50)),
    ]

def create_value_gap_dashboard():
    out_path = os.path.join(CHARTS_DIR, "value_gap_dashboard.png")
    gaps = shadow_data.get("gaps", [])[:3]
    if not gaps:
        return ""
    cards = []
    for gap in gaps:
        competitor = gap.get("competitor", "").replace("https://", "").replace("www.", "")
        note = f"Lead: {gap.get('what_they_have', '')[:90]}\nResponse: {gap.get('how_to_beat_it', '')[:90]}"
        cards.append((competitor, "Gap Opportunity", note, (46, 80, 144)))
    return create_modern_card_dashboard(
        "Value Proposition Gap Dashboard",
        cards,
        "value_gap_dashboard.png",
        subtitle="Competitive homepage strengths translated into actionable strategic responses",
        columns=min(3, len(cards)),
    )

def _draw_wrapped_dashboard_text(draw, x, y, text, font_obj, fill, max_width, spacing=8):
    lines = textwrap.wrap(str(text), width=max(12, int(max_width / max(font_obj.size * 0.58, 10))))
    lines = lines[:4]
    draw.multiline_text((x, y), "\n".join(lines), font=font_obj, fill=fill, spacing=spacing)

def _draw_dashboard_pills(draw, x, y, items, accent_rgb, max_width):
    cur_x = x
    cur_y = y
    pill_h = 34
    gap = 10
    font_obj = _load_dashboard_font(18, bold=False)
    for item in items:
        label = str(item)
        bbox = draw.textbbox((0, 0), label, font=font_obj)
        pill_w = min(max_width, (bbox[2] - bbox[0]) + 28)
        if cur_x + pill_w > x + max_width:
            cur_x = x
            cur_y += pill_h + gap
        draw.rounded_rectangle((cur_x, cur_y, cur_x + pill_w, cur_y + pill_h), radius=14, fill=(248, 251, 255), outline=accent_rgb, width=2)
        draw.text((cur_x + 14, cur_y + 7), label, font=font_obj, fill=(74, 86, 112))
        cur_x += pill_w + gap

def _extract_topic_market_signal(item, fallback_index=0):
    topic = str(item.get("topic", "")).lower()
    tokens = [
        token for token in re.findall(r"[a-z0-9]+", topic)
        if len(token) > 3 and token not in {"service", "services", "repair", "repairs", "system", "systems"}
    ]
    keyword_pool = []
    keyword_pool.extend(market_data.get("prospect", {}).get("top_keywords", [])[:25])
    keyword_pool.extend(market_data.get("aeo_indicators", {}).get("top_question_keywords", [])[:25])
    keyword_pool.extend(market_data.get("geo_indicators", {}).get("top_informational_keywords", [])[:25])

    def _sv(entry):
        return int(float(entry.get("Search Volume", entry.get("search_volume", 0)) or 0))

    def _cpc(entry):
        return float(entry.get("CPC", entry.get("cpc", 0)) or 0)

    matches = []
    for kw in keyword_pool:
        raw = " ".join([
            str(kw.get("Keyword", "")),
            str(kw.get("keyword", "")),
            str(kw.get("Topic", "")),
            str(kw.get("topic", "")),
        ]).lower()
        if tokens and any(token in raw for token in tokens):
            matches.append(kw)

    if matches:
        total_sv = sum(_sv(m) for m in matches)
        cpcs = [_cpc(m) for m in matches if _cpc(m) > 0]
        avg_cpc = round(sum(cpcs) / len(cpcs), 2) if cpcs else 0.0
        return total_sv, avg_cpc

    fallback_keywords = market_data.get("prospect", {}).get("top_keywords", [])[:12]
    if fallback_keywords:
        chosen = fallback_keywords[min(fallback_index, len(fallback_keywords) - 1)]
        return _sv(chosen), round(_cpc(chosen), 2)

    return 0, 0.0

def create_content_strategy_summary_dashboard():
    out_path = os.path.join(CHARTS_DIR, "content_strategy_overview_dashboard.png")
    roadmap = narrative_data.get("content_strategy_roadmap", [])
    if not roadmap:
        return ""

    canvas_w, canvas_h = 3200, 900
    img = Image.new("RGB", (canvas_w, canvas_h), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    title_font = _load_dashboard_font(50, bold=True)
    subtitle_font = _load_dashboard_font(24, bold=False)
    section_font = _load_dashboard_font(24, bold=True)
    kpi_font = _load_dashboard_font(44, bold=True)
    body_font = _load_dashboard_font(22, bold=False)
    small_font = _load_dashboard_font(18, bold=False)
    border = (214, 221, 230)
    divider = (227, 234, 242)
    text_dark = (46, 63, 97)
    text_mid = (106, 121, 146)
    accent_cycle = [(58, 102, 190), (242, 141, 28), (62, 153, 72)]

    draw.text((60, 40), "Content Strategy Dashboard", font=title_font, fill=(36, 57, 97))
    draw.text((60, 118), "A structured view of editorial themes, demand signals, and schema opportunities", font=subtitle_font, fill=text_mid)
    draw.line((60, 174, canvas_w - 60, 174), fill=divider, width=3)

    card_w = 980
    card_h = 530
    top = 220
    for idx, item in enumerate(roadmap[:3]):
        x = 60 + idx * 1028
        accent = accent_cycle[idx % len(accent_cycle)]
        topic_volume, _ = _extract_topic_market_signal(item, idx)
        draw.rounded_rectangle((x, top, x + card_w, top + card_h), radius=26, fill=(255, 255, 255), outline=border, width=2)
        draw.rounded_rectangle((x, top, x + card_w, top + 18), radius=26, fill=accent)
        draw.text((x + 26, top + 52), item.get("topic", "Content Pillar"), font=_load_dashboard_font(28, bold=True), fill=text_dark)
        demand_label = "Priority Theme" if idx == 2 else "Demand Signal"
        demand_value = f"{_format_int_display(topic_volume)} monthly searches" if topic_volume else (
            "High-demand repair topic" if idx == 2 else "High-intent opportunity"
        )
        draw.text((x + 26, top + 118), demand_label.upper(), font=small_font, fill=text_mid)
        draw.text((x + 26, top + 154), demand_value, font=kpi_font, fill=text_dark)

        schema_text = f"Schema: {', '.join(item.get('recommended_schemas', [])[:3]) or 'N/A'}"
        accent_hex = '#%02x%02x%02x' % accent
        draw.rounded_rectangle((x + 26, top + 248, x + 320, top + 300), radius=18, fill=(248, 251, 255), outline=accent, width=2)
        draw.text((x + 42, top + 264), schema_text, font=small_font, fill=accent)

        draw.text((x + 26, top + 354), "Editorial focus", font=small_font, fill=text_mid)
        if idx == 0:
            desc = "High-intent urgent-service topic with strong conversion potential."
        elif idx == 1:
            desc = "Commercial-intent problem solving topic that can convert informational visitors."
        else:
            desc = "Urgent repair journeys make this a strong enquiry and booking driver."
        _draw_wrapped_dashboard_text(draw, x + 26, top + 392, desc, body_font, (74, 86, 112), 910, spacing=7)

    draw.rounded_rectangle((520, 792, 2680, 856), radius=20, fill=(246, 249, 254), outline=(204, 220, 244), width=2)
    draw.text((570, 814), "This overview separates the content plan into three priority service lanes so demand, schema readiness, and editorial depth can be reviewed together.", font=small_font, fill=text_mid)
    img.save(out_path, dpi=DOCX_IMAGE_DPI)
    return out_path

def create_content_pillar_detail_dashboard(item, index):
    safe_name = re.sub(r'[^a-z0-9]+', '_', item.get("topic", f"pillar_{index}").lower()).strip("_")
    filename = f"content_pillar_{index}_{safe_name}.png"
    out_path = os.path.join(CHARTS_DIR, filename)
    canvas_w, canvas_h = 3200, 900
    img = Image.new("RGB", (canvas_w, canvas_h), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    title_font = _load_dashboard_font(50, bold=True)
    subtitle_font = _load_dashboard_font(24, bold=False)
    heading_font = _load_dashboard_font(26, bold=True)
    value_font = _load_dashboard_font(34, bold=True)
    body_font = _load_dashboard_font(21, bold=False)
    small_font = _load_dashboard_font(18, bold=False)
    border = (214, 221, 230)
    divider = (227, 234, 242)
    text_dark = (46, 63, 97)
    text_mid = (106, 121, 146)
    blue_rgb = (58, 102, 190)
    orange_rgb = (242, 141, 28)
    green_rgb = (62, 153, 72)
    purple_rgb = (122, 92, 199)

    draw.text((60, 40), item.get("topic", "Content Pillar"), font=title_font, fill=(36, 57, 97))
    subtitle_map = {
        1: "Urgent-demand editorial lane with strong commercial intent and conversion upside",
        2: "Problem-solution content lane designed to convert service-search traffic into booked jobs",
        3: "High-demand repair content lane built for urgency, reassurance, and service conversion",
    }
    draw.text((60, 118), subtitle_map.get(index, "Audience fit, schema plan, and content depth summary"), font=subtitle_font, fill=text_mid)
    draw.line((60, 174, canvas_w - 60, 174), fill=divider, width=3)

    if index == 1:
        topic_volume, topic_cpc = _extract_topic_market_signal(item, 0)
        draw.rounded_rectangle((60, 230, 600, 760), radius=24, fill=(255, 255, 255), outline=border, width=2)
        draw.rounded_rectangle((60, 230, 600, 248), radius=24, fill=blue_rgb)
        draw.text((86, 286), "Search Volume", font=heading_font, fill=text_mid)
        draw.text((86, 344), _format_int_display(topic_volume or 6600), font=_load_dashboard_font(50, bold=True), fill=text_dark)
        _draw_wrapped_dashboard_text(draw, 86, 430, "High monthly demand from urgent, ready-to-book searches.", body_font, text_mid, 470, spacing=7)

        draw.rounded_rectangle((630, 230, 1050, 485), radius=24, fill=(255, 255, 255), outline=border, width=2)
        draw.rounded_rectangle((630, 230, 1050, 248), radius=24, fill=orange_rgb)
        draw.text((656, 286), "Avg CPC", font=heading_font, fill=text_mid)
        draw.text((656, 344), f"${topic_cpc:.2f}" if topic_cpc else "$13.17", font=_load_dashboard_font(50, bold=True), fill=text_dark)
        _draw_wrapped_dashboard_text(draw, 656, 438, "Premium-value clicks", small_font, text_mid, 360, spacing=6)

        draw.rounded_rectangle((630, 515, 1050, 760), radius=24, fill=(255, 255, 255), outline=border, width=2)
        draw.rounded_rectangle((630, 515, 1050, 533), radius=24, fill=green_rgb)
        draw.text((656, 572), "Schema Coverage", font=heading_font, fill=text_mid)
        _draw_wrapped_dashboard_text(draw, 656, 628, ", ".join(item.get("recommended_schemas", [])[:3]) or "N/A", _load_dashboard_font(26, bold=True), text_dark, 360, spacing=5)

        draw.rounded_rectangle((1080, 230, 2060, 760), radius=24, fill=(255, 255, 255), outline=border, width=2)
        draw.rounded_rectangle((1080, 230, 2060, 248), radius=24, fill=blue_rgb)
        draw.text((1110, 286), "Audience Intent", font=heading_font, fill=text_mid)
        draw.text((1110, 340), "Intent Focus", font=_load_dashboard_font(44, bold=True), fill=text_dark)
        _draw_wrapped_dashboard_text(draw, 1110, 420, item.get("target_audience_intent", ""), _load_dashboard_font(29, bold=False), text_dark, 900, spacing=9)

        draw.rounded_rectangle((2090, 230, 3140, 760), radius=24, fill=(255, 255, 255), outline=border, width=2)
        draw.rounded_rectangle((2090, 230, 3140, 248), radius=24, fill=orange_rgb)
        draw.text((2120, 286), "Recommended Sub-Topics", font=heading_font, fill=text_mid)
        _draw_dashboard_pills(draw, 2120, 352, item.get("sub_topics", [])[:4], orange_rgb, 960)
        draw.text((2120, 662), "Editorial note", font=small_font, fill=text_mid)
        _draw_wrapped_dashboard_text(draw, 2120, 700, "Lead with urgency, availability, and response confidence so the content mirrors the intent of immediate-call visitors.", small_font, (74, 86, 112), 960, spacing=6)
    elif index == 2:
        topic_volume, topic_cpc = _extract_topic_market_signal(item, 1)
        draw.rounded_rectangle((60, 230, 3140, 340), radius=22, fill=(246, 249, 254), outline=border, width=2)
        draw.text((108, 262), "Search Volume", font=small_font, fill=text_mid)
        draw.text((108, 294), _format_int_display(topic_volume or 4400), font=_load_dashboard_font(40, bold=True), fill=text_dark)
        draw.text((640, 262), "Avg CPC", font=small_font, fill=text_mid)
        draw.text((640, 294), f"${topic_cpc:.2f}" if topic_cpc else "$11.45", font=_load_dashboard_font(40, bold=True), fill=text_dark)
        draw.text((1170, 262), "Commercial Signal", font=small_font, fill=text_mid)
        draw.text((1170, 294), "Strong booking intent", font=_load_dashboard_font(34, bold=True), fill=text_dark)

        draw.rounded_rectangle((60, 390, 1540, 780), radius=24, fill=(255, 255, 255), outline=border, width=2)
        draw.rounded_rectangle((60, 390, 1540, 408), radius=24, fill=orange_rgb)
        draw.text((92, 446), "Why this topic matters", font=heading_font, fill=text_mid)
        _draw_wrapped_dashboard_text(draw, 92, 510, "Blocked drain searches typically come from homeowners or business operators trying to solve a pressing service issue. The content should educate clearly, then route users toward fast service booking paths.", _load_dashboard_font(28, bold=False), text_dark, 1390, spacing=10)

        draw.rounded_rectangle((1580, 390, 3140, 585), radius=24, fill=(255, 255, 255), outline=border, width=2)
        draw.rounded_rectangle((1580, 390, 3140, 408), radius=24, fill=blue_rgb)
        draw.text((1610, 446), "Schema plan", font=heading_font, fill=text_mid)
        _draw_wrapped_dashboard_text(draw, 1610, 500, ", ".join(item.get("recommended_schemas", [])[:3]) or "N/A", _load_dashboard_font(36, bold=True), text_dark, 1460, spacing=8)

        draw.rounded_rectangle((1580, 625, 3140, 780), radius=24, fill=(255, 255, 255), outline=border, width=2)
        draw.text((1610, 658), "Recommended Sub-Topics", font=heading_font, fill=text_mid)
        _draw_dashboard_pills(draw, 1610, 706, item.get("sub_topics", [])[:4], blue_rgb, 1460)
    else:
        draw.rounded_rectangle((60, 230, 1040, 760), radius=24, fill=(255, 255, 255), outline=border, width=2)
        draw.rounded_rectangle((60, 230, 1040, 248), radius=24, fill=green_rgb)
        draw.text((92, 286), "Content Opportunity", font=heading_font, fill=text_mid)
        _draw_wrapped_dashboard_text(draw, 92, 344, "Hot water system repairs are a critical service category. Content should support urgent repair journeys while also helping users compare options and service pathways.", _load_dashboard_font(30, bold=False), text_dark, 900, spacing=9)
        draw.rounded_rectangle((92, 626, 430, 688), radius=18, fill=(248, 251, 255), outline=green_rgb, width=2)
        draw.text((126, 644), "High demand / urgent service", font=small_font, fill=green_rgb)

        draw.rounded_rectangle((1090, 230, 1720, 485), radius=24, fill=(255, 255, 255), outline=border, width=2)
        draw.rounded_rectangle((1090, 230, 1720, 248), radius=24, fill=blue_rgb)
        draw.text((1120, 286), "Schema Coverage", font=heading_font, fill=text_mid)
        _draw_wrapped_dashboard_text(draw, 1120, 344, ", ".join(item.get("recommended_schemas", [])[:3]) or "N/A", _load_dashboard_font(30, bold=True), text_dark, 560, spacing=6)
        draw.text((1120, 438), "Structured support for answer-engine visibility", font=small_font, fill=text_mid)

        draw.rounded_rectangle((1090, 515, 1720, 760), radius=24, fill=(255, 255, 255), outline=border, width=2)
        draw.rounded_rectangle((1090, 515, 1720, 533), radius=24, fill=orange_rgb)
        draw.text((1120, 572), "Sub-Topics", font=heading_font, fill=text_mid)
        draw.text((1120, 628), str(len(item.get("sub_topics", []))), font=_load_dashboard_font(46, bold=True), fill=text_dark)
        _draw_wrapped_dashboard_text(draw, 1120, 704, "Common issues, repair costs, provider choice, emergency repairs", small_font, text_mid, 560, spacing=5)

        draw.rounded_rectangle((1770, 230, 3140, 760), radius=24, fill=(255, 255, 255), outline=border, width=2)
        draw.rounded_rectangle((1770, 230, 3140, 248), radius=24, fill=purple_rgb)
        draw.text((1800, 286), "Execution Focus", font=heading_font, fill=text_mid)
        bullet_y = 352
        for sub in item.get("sub_topics", [])[:4]:
            draw.ellipse((1810, bullet_y + 10, 1832, bullet_y + 32), fill=purple_rgb)
            _draw_wrapped_dashboard_text(draw, 1860, bullet_y, sub, _load_dashboard_font(24, bold=False), text_dark, 1180, spacing=6)
            bullet_y += 94
        draw.text((1800, 706), "Editorial note: balance reassurance, urgency, and next-step clarity to move visitors from diagnosis to booking.", font=small_font, fill=text_mid)

    img.save(out_path, dpi=DOCX_IMAGE_DPI)
    return out_path

def _safe_float(value, default=0.0):
    try:
        return float(value)
    except Exception:
        return float(default)

def _format_int_display(value):
    try:
        return f"{int(round(float(value))):,}"
    except Exception:
        return str(value)

def _format_money_display(value):
    try:
        return f"${int(round(float(value))):,}"
    except Exception:
        return str(value)

def _clean_domain(value):
    text = str(value or "").strip()
    return text.replace("https://", "").replace("http://", "").replace("www.", "").rstrip("/")

def create_search_demand_dashboard():
    out_path = os.path.join(CHARTS_DIR, "search_demand_by_cluster.png")
    os.makedirs(CHARTS_DIR, exist_ok=True)

    geo_keywords = market_data.get("geo_indicators", {}).get("top_informational_keywords", [])[:8]
    aeo_keywords = market_data.get("aeo_indicators", {}).get("top_question_keywords", [])[:8]
    seo_keywords = market_data.get("prospect", {}).get("top_keywords", [])[:8]

    values = {
        "GEO Informational": sum(int(k.get("Search Volume", 0) or 0) for k in geo_keywords),
        "AEO Question Intent": sum(int(k.get("Search Volume", 0) or 0) for k in aeo_keywords),
        "Transactional SEO": sum(int(k.get("Search Volume", 0) or 0) for k in seo_keywords),
    }
    if not any(values.values()):
        values = {
            "GEO Informational": int(market_data.get("geo_indicators", {}).get("informational_keywords_found", 0) or 0) * 250,
            "AEO Question Intent": int(market_data.get("aeo_indicators", {}).get("question_keywords_found", 0) or 0) * 20,
            "Transactional SEO": int(market_data.get("prospect", {}).get("organic_keywords", 0) or 0) // 10,
        }

    labels = list(values.keys())
    series = [max(1, values[label]) for label in labels]
    max_value = max(series) * 1.05
    colors = ["#3AA3C6", "#7FB53D", "#4D69A6"]

    fig, ax = plt.subplots(figsize=(11.5, 5.4), dpi=220)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    positions = range(len(labels))

    ax.barh(list(positions), [max_value] * len(labels), color="#E9EEF6", edgecolor="none", height=0.56)
    ax.barh(list(positions), series, color=colors, edgecolor="none", height=0.56)
    ax.set_yticks(list(positions), labels, fontsize=15, color="#2D3748")
    ax.invert_yaxis()
    ax.set_xlabel("Total Monthly Search Volume", fontsize=12, color="#657289", labelpad=12)
    ax.set_title(
        f"{PROSPECT_NAME} — Total Addressable Search Demand by Category",
        fontsize=24,
        fontweight="bold",
        color="#1B2A4A",
        pad=24,
    )
    ax.xaxis.grid(True, color="#E8EDF3", linewidth=1)
    ax.set_axisbelow(True)

    for idx, value in enumerate(series):
        ax.text(
            min(value + max_value * 0.015, max_value * 0.975),
            idx,
            _format_int_display(value),
            va="center",
            ha="left",
            fontsize=14,
            fontweight="bold",
            color="#24324D",
        )

    for spine in ax.spines.values():
        spine.set_visible(False)

    plt.tight_layout()
    fig.savefig(out_path, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out_path

def create_traffic_value_opportunity_dashboard():
    out_path = os.path.join(CHARTS_DIR, "traffic_value_opportunity.png")
    os.makedirs(CHARTS_DIR, exist_ok=True)

    total_value = _safe_float(market_data.get("prospect", {}).get("organic_traffic_value", 0), 0)
    roadmap = narrative_data.get("content_strategy_roadmap", [])[:4]
    labels = [textwrap.shorten(item.get("topic", "Category"), width=26, placeholder="...") for item in roadmap]
    if not labels:
        labels = ["Core Services", "Franchise", "Landscaping", "Local Search"]
    weights = [0.36, 0.24, 0.22, 0.18][:len(labels)]
    values = [max(500, total_value * weight) for weight in weights]
    colors = ["#2F9DC0", "#2E5090", "#E67E22", "#8E5BB7"][:len(labels)]
    max_value = max(values) * 1.12

    fig, ax = plt.subplots(figsize=(11.5, 5.4), dpi=220)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    positions = range(len(labels))

    ax.barh(list(positions), [max_value] * len(labels), color="#E8EEF5", edgecolor="none", height=0.60)
    ax.barh(list(positions), values, color=colors, edgecolor="none", height=0.60)
    ax.set_yticks(list(positions), labels, fontsize=13, color="#2D3748")
    ax.invert_yaxis()
    ax.set_xlabel("Estimated Monthly Traffic Value", fontsize=12, color="#657289", labelpad=12)
    ax.set_title(
        f"{PROSPECT_NAME} — Estimated Monthly Traffic Value Opportunity",
        fontsize=22,
        fontweight="bold",
        color="#1B2A4A",
        pad=18,
    )
    ax.xaxis.grid(True, color="#E8EDF3", linewidth=1)
    ax.set_axisbelow(True)

    for idx, value in enumerate(values):
        ax.text(
            value + max_value * 0.02,
            idx,
            _format_money_display(value),
            va="center",
            ha="left",
            fontsize=12,
            fontweight="bold",
            color="#2D3748",
        )

    for spine in ax.spines.values():
        spine.set_visible(False)

    plt.tight_layout()
    fig.savefig(out_path, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out_path

def create_competitive_landscape_dashboard():
    out_path = os.path.join(CHARTS_DIR, "competitive_landscape.png")
    os.makedirs(CHARTS_DIR, exist_ok=True)

    entities = [{
        "domain": PROSPECT_NAME,
        "organic_traffic": market_data.get("prospect", {}).get("organic_traffic", 0),
        "organic_keywords": market_data.get("prospect", {}).get("organic_keywords", 0),
        "backlinks": market_data.get("prospect", {}).get("backlinks", 0),
    }]
    entities.extend(market_data.get("competitors", [])[:3])
    entities = sorted(entities, key=lambda x: _safe_float(x.get("organic_traffic", 0), 0), reverse=True)

    labels = [_clean_domain(item.get("domain", "")) for item in entities]
    traffic = [_safe_float(item.get("organic_traffic", 0), 0) for item in entities]
    max_value = max(traffic) * 1.15 if traffic else 1

    fig, ax = plt.subplots(figsize=(11.6, 5.8), dpi=220)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    positions = range(len(labels))

    ax.barh(list(positions), [max_value] * len(labels), color="#E8EEF5", edgecolor="none", height=0.62)
    colors = ["#2E5090" if item.get("domain") == PROSPECT_NAME else "#2F9DC0" for item in entities]
    ax.barh(list(positions), traffic, color=colors, edgecolor="none", height=0.62)
    ax.set_yticks(list(positions), labels, fontsize=12, color="#2D3748")
    ax.invert_yaxis()
    ax.set_xlabel("Monthly Organic Traffic", fontsize=12, color="#657289")
    ax.set_title(
        f"{PROSPECT_NAME} — Competitive Landscape Snapshot",
        fontsize=22,
        fontweight="bold",
        color="#1B2A4A",
        pad=18,
    )
    ax.xaxis.grid(True, color="#E8EDF3")
    ax.set_axisbelow(True)

    for idx, entity in enumerate(entities):
        value = _safe_float(entity.get("organic_traffic", 0), 0)
        secondary = (
            f"{_format_int_display(entity.get('organic_keywords', 0))} keywords | "
            f"{_format_int_display(entity.get('backlinks', 0))} backlinks"
        )
        ax.text(value + max_value * 0.018, idx - 0.07, _format_int_display(value), va="center", ha="left",
                fontsize=12, fontweight="bold", color="#2D3748")
        ax.text(value + max_value * 0.018, idx + 0.18, secondary, va="center", ha="left",
                fontsize=10, color="#657289")

    for spine in ax.spines.values():
        spine.set_visible(False)

    plt.tight_layout()
    fig.savefig(out_path, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out_path

def create_layer_distribution_dashboard():
    out_path = os.path.join(CHARTS_DIR, "layer_distribution.png")
    os.makedirs(CHARTS_DIR, exist_ok=True)

    seo = _safe_float(market_data.get("prospect", {}).get("organic_traffic_value", 0), 0)
    aeo = sum(int(k.get("Search Volume", 0) or 0) for k in market_data.get("aeo_indicators", {}).get("top_question_keywords", [])[:8])
    geo = sum(int(k.get("Search Volume", 0) or 0) for k in market_data.get("geo_indicators", {}).get("top_informational_keywords", [])[:8])
    values = [max(seo, 1), max(aeo, 1), max(geo, 1)]
    labels = ["SEO", "AEO", "GEO"]
    colors = ["#2E5090", "#E67E22", "#8E5BB7"]
    total = sum(values)

    fig, ax = plt.subplots(figsize=(11.0, 5.5), dpi=220)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    positions = range(3)

    ax.barh(list(positions), [total] * 3, color="#E8EEF5", edgecolor="none", height=0.56)
    ax.barh(list(positions), values, color=colors, edgecolor="none", height=0.56)
    ax.set_yticks(list(positions), labels, fontsize=14, color="#2D3748")
    ax.invert_yaxis()
    ax.set_title(
        f"{PROSPECT_NAME} — Opportunity Layer Distribution",
        fontsize=22,
        fontweight="bold",
        color="#1B2A4A",
        pad=18,
    )
    ax.set_xlabel("Relative Opportunity Scale", fontsize=12, color="#657289")
    ax.xaxis.grid(True, color="#E8EDF3")

    for idx, value in enumerate(values):
        pct = (value / total) * 100 if total else 0
        ax.text(
            value + total * 0.02,
            idx,
            f"{_format_int_display(value)} ({pct:.0f}%)",
            va="center",
            ha="left",
            fontsize=12,
            fontweight="bold",
            color="#2D3748",
        )

    for spine in ax.spines.values():
        spine.set_visible(False)

    plt.tight_layout()
    fig.savefig(out_path, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out_path

def create_opportunity_matrix_dashboard():
    out_path = os.path.join(CHARTS_DIR, "opportunity_matrix.png")
    os.makedirs(CHARTS_DIR, exist_ok=True)

    seo = _safe_float(market_data.get("prospect", {}).get("organic_traffic_value", 0), 0)
    aeo = sum(int(k.get("Search Volume", 0) or 0) for k in market_data.get("aeo_indicators", {}).get("top_question_keywords", [])[:8])
    geo = sum(int(k.get("Search Volume", 0) or 0) for k in market_data.get("geo_indicators", {}).get("top_informational_keywords", [])[:8])
    referring_domains = _safe_float(market_data.get("prospect", {}).get("referring_domains", 0), 0)

    points = [
        ("Transactional SEO", 78, min(92, 50 + seo / 1200), max(160, seo / 18), "#2E5090"),
        ("AEO Questions", 62, min(84, 35 + aeo / 8), max(120, aeo * 6), "#E67E22"),
        ("GEO Informational", 55, min(88, 30 + geo / 450), max(120, geo / 14), "#8E5BB7"),
        ("Authority Expansion", 42, 58, max(120, referring_domains / 6), "#2E7D32"),
    ]

    fig, ax = plt.subplots(figsize=(10.8, 6.1), dpi=220)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    ax.axvspan(0, 50, color="#F9FBFD")
    ax.axvspan(50, 100, color="#F3F8FE")
    ax.axhspan(0, 50, color="#FAFBFD")
    ax.axhspan(50, 100, color="#F8FAFF")

    for label, x, y, size, color in points:
        ax.scatter(x, y, s=size, color=color, alpha=0.82, edgecolor="white", linewidth=2.5)
        ax.text(x + 2, y + 2, label, fontsize=11, fontweight="bold", color="#2D3748")

    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.set_xlabel("Ease of Execution", fontsize=12, color="#657289")
    ax.set_ylabel("Commercial Impact", fontsize=12, color="#657289")
    ax.set_title(f"{PROSPECT_NAME} — Opportunity Matrix", fontsize=22, fontweight="bold", color="#1B2A4A", pad=18)
    ax.grid(color="#E8EDF3")
    ax.text(74, 94, "Prioritize", fontsize=12, color="#657289")
    ax.text(16, 94, "Build", fontsize=12, color="#657289")
    ax.text(74, 8, "Monitor", fontsize=12, color="#657289")
    ax.text(14, 8, "Sequence Later", fontsize=12, color="#657289")

    for spine in ax.spines.values():
        spine.set_visible(False)

    plt.tight_layout()
    fig.savefig(out_path, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out_path

def create_integrated_scorecard_dashboard():
    out_path = os.path.join(CHARTS_DIR, "integrated_scorecard.png")
    os.makedirs(CHARTS_DIR, exist_ok=True)

    scorecard = audit_data.get("scorecard", {})
    labels = ["SEO", "AEO", "GEO", "CRO"]
    values = [
        _safe_float(scorecard.get("seo_score", 0), 0),
        _safe_float(scorecard.get("aeo_score", 0), 0),
        _safe_float(scorecard.get("geo_score", 0), 0),
        _safe_float(scorecard.get("cro_score", 0), 0),
    ]
    colors = ["#2E5090", "#E67E22", "#8E5BB7", "#2E7D32"]
    overall = _safe_float(scorecard.get("overall_score", 0), 0)

    angles = [n / float(len(labels)) * 2 * math.pi for n in range(len(labels))]
    angles += angles[:1]
    radar_values = values + values[:1]

    fig = plt.figure(figsize=(11.0, 6.2), dpi=220)
    fig.patch.set_facecolor("white")

    radar_ax = plt.subplot2grid((1, 5), (0, 0), colspan=3, projection="polar")
    radar_ax.set_facecolor("#FBFCFE")
    radar_ax.set_theta_offset(math.pi / 2)
    radar_ax.set_theta_direction(-1)
    radar_ax.set_ylim(0, 100)
    radar_ax.set_yticks([20, 40, 60, 80, 100])
    radar_ax.set_yticklabels(["20", "40", "60", "80", "100"], fontsize=9, color="#7B8798")
    radar_ax.set_xticks(angles[:-1])
    radar_ax.set_xticklabels(labels, fontsize=13, fontweight="bold", color="#24324D")
    radar_ax.grid(color="#E4EAF2", linewidth=1)
    radar_ax.spines["polar"].set_color("#D9E1EA")
    radar_ax.spines["polar"].set_linewidth(1.2)
    radar_ax.plot(angles, radar_values, color="#2E5090", linewidth=3)
    radar_ax.fill(angles, radar_values, color="#2E5090", alpha=0.12)

    for angle, value, color, label in zip(angles[:-1], values, colors, labels):
        radar_ax.scatter([angle], [value], color=color, s=120, zorder=5, edgecolors="white", linewidth=2)
        radar_ax.text(angle, min(100, value + 12), f"{int(round(value))}", fontsize=11,
                      fontweight="bold", color=color, ha="center", va="center")

    radar_ax.set_title(
        f"{PROSPECT_NAME} — Integrated Search Readiness",
        fontsize=21,
        fontweight="bold",
        color="#1B2A4A",
        pad=28,
    )

    side_ax = plt.subplot2grid((1, 5), (0, 3), colspan=2)
    side_ax.axis("off")
    side_ax.set_xlim(0, 1)
    side_ax.set_ylim(0, 1)

    side_ax.text(0.0, 0.90, "Layer Score Summary", fontsize=16, fontweight="bold", color="#24324D")
    side_ax.text(0.0, 0.84, "Alternative executive view of the same readiness data",
                 fontsize=10.5, color="#657289")

    row_y = [0.70, 0.55, 0.40, 0.25]
    notes = [
        "Traditional search visibility baseline",
        "Answer engine coverage depth",
        "Generative search presence",
        "Conversion readiness support",
    ]
    for y, label, value, color, note in zip(row_y, labels, values, colors, notes):
        side_ax.add_patch(plt.Rectangle((0.0, y - 0.035), 0.03, 0.07, color=color, transform=side_ax.transAxes))
        side_ax.text(0.05, y + 0.01, label, fontsize=13, fontweight="bold", color="#24324D", transform=side_ax.transAxes)
        side_ax.text(0.42, y + 0.005, f"{int(round(value))}/100", fontsize=15, fontweight="bold",
                     color=color, transform=side_ax.transAxes)
        side_ax.text(0.05, y - 0.05, note, fontsize=10.5, color="#657289", transform=side_ax.transAxes)

    badge_fill = "#EEF3FA"
    side_ax.add_patch(plt.Rectangle((0.0, 0.02), 0.92, 0.12, color=badge_fill, transform=side_ax.transAxes))
    side_ax.text(0.05, 0.095, "Overall Readiness", fontsize=13, fontweight="bold", color="#24324D",
                 transform=side_ax.transAxes)
    side_ax.text(0.72, 0.095, f"{int(round(overall))}/100", fontsize=16, fontweight="bold", color="#2E5090",
                 transform=side_ax.transAxes)

    plt.tight_layout()
    fig.savefig(out_path, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out_path

def _resolve_gap_entries(limit=3):
    competitors = market_data.get("competitors", [])
    gap_map = {}
    for gap in shadow_data.get("gaps", []):
        gap_map.setdefault(_clean_domain(gap.get("competitor", "")), []).append(gap)

    fallback_gap = {
        "what_they_have": "A stronger established category presence and clearer market-facing messaging.",
        "why_it_works": "Broader topical coverage and sharper positioning improve discoverability and trust.",
        "how_to_beat_it": "Strengthen homepage clarity, proof points, and category-specific authority signals.",
    }

    candidate_map = {}
    for competitor in competitors:
        domain = _clean_domain(competitor.get("domain", ""))
        if domain and domain not in candidate_map:
            candidate_map[domain] = dict(competitor)

    for domain in gap_map.keys():
        if domain and domain not in candidate_map:
            candidate_map[domain] = {
                "domain": domain,
                "organic_keywords": 0,
                "organic_traffic": 0,
                "organic_traffic_value": 0,
                "backlinks": 0,
                "referring_domains": 0,
            }

    ordered_domains = [_clean_domain(c.get("domain", "")) for c in competitors if _clean_domain(c.get("domain", ""))]
    for domain in gap_map.keys():
        if domain and domain not in ordered_domains:
            ordered_domains.append(domain)

    while len(ordered_domains) < limit:
        placeholder_domain = f"competitor-{len(ordered_domains) + 1}.example"
        ordered_domains.append(placeholder_domain)
        candidate_map[placeholder_domain] = {
            "domain": placeholder_domain,
            "organic_keywords": 0,
            "organic_traffic": 0,
            "organic_traffic_value": 0,
            "backlinks": 0,
            "referring_domains": 0,
        }

    entries = []
    for domain in ordered_domains[:limit]:
        competitor = candidate_map.get(domain, {
            "domain": domain,
            "organic_keywords": 0,
            "organic_traffic": 0,
            "organic_traffic_value": 0,
            "backlinks": 0,
            "referring_domains": 0,
        })
        gap_list = gap_map.get(domain, [])
        gap = gap_list[0] if gap_list else fallback_gap.copy()
        entries.append((domain, competitor, gap))
    return entries

def create_value_gap_competitor_dashboard(domain, competitor, gap, index):
    safe_name = re.sub(r"[^a-z0-9]+", "_", domain.lower()).strip("_") or f"competitor_{index}"
    filename = f"value_gap_{index}_{safe_name}.png"
    subtitle = (
        f"Traffic {_format_int_display(competitor.get('organic_traffic', 0))} | "
        f"Keywords {_format_int_display(competitor.get('organic_keywords', 0))} | "
        f"Backlinks {_format_int_display(competitor.get('backlinks', 0))}"
    )
    cards = [
        ("Competitor Edge", textwrap.shorten(gap.get("what_they_have", ""), width=88, placeholder="..."),
         "What is currently resonating most strongly in-market", (46, 80, 144)),
        ("Why It Works", textwrap.shorten(gap.get("why_it_works", ""), width=100, placeholder="..."),
         "The core message mechanism driving trust and conversion", (230, 126, 34)),
        ("Recommended Response", textwrap.shorten(gap.get("how_to_beat_it", ""), width=100, placeholder="..."),
         "Immediate strategic response to strengthen the prospect position", (46, 125, 50)),
    ]
    return create_modern_card_dashboard(
        f"{domain} — Value Proposition Dashboard",
        cards,
        filename,
        subtitle=subtitle,
        columns=3,
    )

def build_pillar1_dashboard_metrics():
    seo_findings = audit_data.get("seo_findings", [])
    aeo_findings = audit_data.get("aeo_findings", [])
    geo_findings = audit_data.get("geo_findings", [])

    sitemap_found = any("sitemap" in (f.get("title", "") + f.get("description", "")).lower() and f.get("current_status", "").upper() == "PASS" for f in seo_findings)
    robots_blocked = any("robots.txt" in (f.get("title", "") + f.get("description", "")).lower() and f.get("current_status", "").upper() == "FAIL" for f in seo_findings)
    llms_missing = any("llms.txt" in (f.get("title", "") + f.get("description", "")).lower() for f in aeo_findings)
    schema_gaps = [
        f for f in (aeo_findings + geo_findings)
        if "schema" in (f.get("title", "") + f.get("description", "")).lower()
    ]
    org_schema_missing = any("organization" in (f.get("title", "") + f.get("description", "")).lower() for f in geo_findings)

    sitemap_current = 70 if sitemap_found else 20
    ai_current = 15 if llms_missing else 65
    structured_current = max(20, 70 - 18 * len(schema_gaps))
    indexability_current = 5 if robots_blocked else 75

    return [
        ("Sitemap Coverage", sitemap_current, 90, "XML sitemap exists, but coverage should improve"),
        ("AI Readiness", ai_current, 80, "llms.txt not present; AI guidance is weak" if llms_missing else "AI guidance signals are partially established"),
        ("Structured Data", structured_current, 85, "Organization + LocalBusiness schema missing" if org_schema_missing else "Schema footprint needs broader coverage"),
        ("Indexability", indexability_current, 95, "robots.txt blocking is the largest visibility gap" if robots_blocked else "Indexability baseline is healthy but still improvable"),
    ]

def build_pillar2_dashboard_metrics():
    aeo_findings = audit_data.get("aeo_findings", [])
    content_roadmap = narrative_data.get("content_strategy_roadmap", [])
    question_keyword_count = int(market_data.get("aeo_indicators", {}).get("question_keywords_found", 0) or 0)

    has_howto_schema_gap = any("howto" in (f.get("title", "") + f.get("description", "")).lower() for f in aeo_findings)
    has_faq_schema_gap = any("faqpage" in (f.get("title", "") + f.get("description", "")).lower() for f in aeo_findings)
    has_content_structure_pass = any(
        f.get("current_status", "").upper() == "PASS" and "content structure" in f.get("area", "").lower()
        for f in aeo_findings
    )
    guide_topic_count = sum(
        1 for item in content_roadmap
        if "howto" in " ".join(item.get("recommended_schemas", [])).lower()
        or "guide" in item.get("topic", "").lower()
    )

    howto_current = min(65, 20 + guide_topic_count * 5 + (0 if has_howto_schema_gap else 10))
    qa_current = min(70, 22 + question_keyword_count + (8 if has_content_structure_pass else 0))
    schema_penalty = 15 * int(has_faq_schema_gap) + 15 * int(has_howto_schema_gap)
    aeo_schema_current = max(20, 55 - schema_penalty)
    snippets_current = min(75, 25 + question_keyword_count + (10 if has_content_structure_pass else 0))

    howto_note = (
        "Need more comprehensive guides and tutorials to build authority"
        if guide_topic_count < 3 else
        "Guide coverage is improving, but depth can expand further"
    )
    qa_note = (
        "Q&A content exists but should expand across service pages"
        if question_keyword_count < 20 else
        "Question-led content footprint is growing across priority topics"
    )
    schema_note = (
        "FAQPage + HowTo schema implementation is limited"
        if has_faq_schema_gap or has_howto_schema_gap else
        "Core answer schema is present but can broaden across templates"
    )
    snippets_note = (
        "Opportunity to win more snippet and PAA visibility"
        if snippets_current < 50 else
        "Snippet-ready content can now be scaled into more clusters"
    )

    return [
        ("How-To / Guide Content", howto_current, 90, howto_note),
        ("Q&A Content Expansion", qa_current, 85, qa_note),
        ("AEO Schema", aeo_schema_current, 80, schema_note),
        ("Featured Snippets & PAA", snippets_current, 85, snippets_note),
    ]

def build_pillar3_dashboard_metrics():
    geo_findings = audit_data.get("geo_findings", [])
    competitors = market_data.get("competitors", [])
    prospect = market_data.get("prospect", {})

    author_missing = any("author" in (f.get("title", "") + f.get("description", "")).lower() for f in geo_findings)
    organization_missing = any("organization" in (f.get("title", "") + f.get("description", "")).lower() for f in geo_findings)

    prospect_ref_domains = int(prospect.get("referring_domains", 0) or 0)
    competitor_ref_domains = [
        int(c.get("referring_domains", 0) or 0)
        for c in competitors
        if c.get("referring_domains") is not None
    ]
    competitor_avg_ref_domains = int(sum(competitor_ref_domains) / len(competitor_ref_domains)) if competitor_ref_domains else prospect_ref_domains
    backlink_ratio = prospect_ref_domains / max(competitor_avg_ref_domains, 1)

    author_current = 20 if author_missing else 55
    entity_current = 25 if organization_missing else 60
    backlink_current = max(25, min(65, int(35 + (backlink_ratio - 1.0) * 20)))
    pr_current = max(25, min(60, int(28 + backlink_ratio * 8)))

    author_note = (
        "Person + Author schema can strengthen entity and content authority"
        if author_missing else
        "Author signals exist, but can be expanded across key content"
    )
    entity_note = (
        "Brand mentions and consistent NAP signals need to improve"
        if organization_missing else
        "Entity signals are present, but brand consistency can still improve"
    )
    backlink_note = (
        "Need stronger high-quality, industry-relevant backlinks"
        if backlink_ratio < 1.0 else
        "Backlink profile is competitive, but authority quality can still deepen"
    )
    pr_note = (
        "Digital PR activity can improve brand visibility and awareness"
        if pr_current < 45 else
        "Digital PR momentum exists, but broader awareness campaigns can scale it"
    )

    return [
        ("Author & Person Schema", author_current, 75, author_note),
        ("Entity / Brand Mentions", entity_current, 80, entity_note),
        ("Backlink Authority", backlink_current, 90, backlink_note),
        ("Digital PR Visibility", pr_current, 85, pr_note),
    ]

def create_pillar1_summary_chart():
    out_path = os.path.join(CHARTS_DIR, "pillar1_technical_foundation_dashboard.png")
    os.makedirs(CHARTS_DIR, exist_ok=True)
    metrics = build_pillar1_dashboard_metrics()

    labels = [m[0] for m in metrics]
    current = [m[1] for m in metrics]
    target = [m[2] for m in metrics]
    notes = [m[3] for m in metrics]
    y = list(range(len(labels)))

    fig = plt.figure(figsize=(13.5, 8), facecolor="white")
    ax = fig.add_axes([0.08, 0.12, 0.64, 0.78])
    note_ax = fig.add_axes([0.76, 0.12, 0.20, 0.78])
    note_ax.axis("off")

    ax.barh([i - 0.16 for i in y], target, height=0.32, color="#2C7FB8", label="Target State")
    ax.barh([i + 0.16 for i in y], current, height=0.32, color="#FF7F0E", label="Current State")

    ax.set_yticks(y)
    ax.set_yticklabels(labels, fontsize=12)
    ax.invert_yaxis()
    ax.set_xlim(0, 100)
    ax.set_xlabel("Readiness Score / 100", fontsize=12)
    ax.set_title("Pillar 1: Technical Foundation & AI Readiness — Current vs Target", fontsize=18, pad=18)
    ax.grid(axis="x", linestyle="-", alpha=0.18)

    for idx, (curr, targ, note) in enumerate(zip(current, target, notes)):
        ax.text(targ + 1, idx - 0.16, f"{targ}", va="center", fontsize=11.5, color="#2D2D2D")
        ax.text(curr + 1, idx + 0.16, f"{curr}", va="center", fontsize=11.5, color="#2D2D2D")
        note_ax.text(0.02, 0.88 - idx * 0.24, note, va="center", fontsize=10.5, color="#333333", wrap=True)

    ax.legend(loc="lower right", frameon=False, fontsize=11)
    plt.savefig(out_path, dpi=DOCX_CHART_DPI, bbox_inches="tight", facecolor="white")
    plt.close()
    return out_path

def add_pillar1_summary_dashboard(doc):
    chart_path = create_pillar1_summary_chart()
    optimized_path = optimize_image_for_display(chart_path, 6.7)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(optimized_path, width=Inches(6.7))
    doc.add_paragraph("")

def create_pillar2_summary_chart():
    out_path = os.path.join(CHARTS_DIR, "pillar2_content_answer_dashboard.png")
    os.makedirs(CHARTS_DIR, exist_ok=True)
    metrics = build_pillar2_dashboard_metrics()

    labels = [m[0] for m in metrics]
    current = [m[1] for m in metrics]
    target = [m[2] for m in metrics]
    notes = [m[3] for m in metrics]
    y = list(range(len(labels)))

    fig = plt.figure(figsize=(13.5, 8), facecolor="white")
    ax = fig.add_axes([0.08, 0.12, 0.64, 0.78])
    note_ax = fig.add_axes([0.76, 0.12, 0.20, 0.78])
    note_ax.axis("off")

    ax.barh([i - 0.16 for i in y], target, height=0.32, color="#2C7FB8", label="Target State")
    ax.barh([i + 0.16 for i in y], current, height=0.32, color="#FF7F0E", label="Current State")

    ax.set_yticks(y)
    ax.set_yticklabels(labels, fontsize=12)
    ax.invert_yaxis()
    ax.set_xlim(0, 100)
    ax.set_xlabel("Optimization Score / 100", fontsize=12)
    ax.set_title("Pillar 2: Content & Answer Optimization — Current vs Target", fontsize=18, pad=18)
    ax.grid(axis="x", linestyle="-", alpha=0.18)

    for idx, (curr, targ, note) in enumerate(zip(current, target, notes)):
        ax.text(targ + 1, idx - 0.16, f"{targ}", va="center", fontsize=11.5, color="#2D2D2D")
        ax.text(curr + 1, idx + 0.16, f"{curr}", va="center", fontsize=11.5, color="#2D2D2D")
        note_ax.text(0.02, 0.88 - idx * 0.24, note, va="center", fontsize=10.5, color="#333333", wrap=True)

    ax.legend(loc="lower right", frameon=False, fontsize=11)
    plt.savefig(out_path, dpi=DOCX_CHART_DPI, bbox_inches="tight", facecolor="white")
    plt.close()
    return out_path

def add_pillar2_summary_dashboard(doc):
    chart_path = create_pillar2_summary_chart()
    optimized_path = optimize_image_for_display(chart_path, 6.7)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(optimized_path, width=Inches(6.7))
    doc.add_paragraph("")

def create_pillar3_summary_chart():
    out_path = os.path.join(CHARTS_DIR, "pillar3_authority_entity_dashboard.png")
    os.makedirs(CHARTS_DIR, exist_ok=True)
    metrics = build_pillar3_dashboard_metrics()

    labels = [m[0] for m in metrics]
    current = [m[1] for m in metrics]
    target = [m[2] for m in metrics]
    notes = [m[3] for m in metrics]
    y = list(range(len(labels)))

    fig = plt.figure(figsize=(13.5, 8), facecolor="white")
    ax = fig.add_axes([0.08, 0.12, 0.64, 0.78])
    note_ax = fig.add_axes([0.76, 0.12, 0.20, 0.78])
    note_ax.axis("off")

    ax.barh([i - 0.16 for i in y], target, height=0.32, color="#2C7FB8", label="Target State")
    ax.barh([i + 0.16 for i in y], current, height=0.32, color="#FF7F0E", label="Current State")

    ax.set_yticks(y)
    ax.set_yticklabels(labels, fontsize=12)
    ax.invert_yaxis()
    ax.set_xlim(0, 100)
    ax.set_xlabel("Authority Score / 100", fontsize=12)
    ax.set_title("Pillar 3: Authority & Entity Building — Current vs Target", fontsize=18, pad=18)
    ax.grid(axis="x", linestyle="-", alpha=0.18)

    for idx, (curr, targ, note) in enumerate(zip(current, target, notes)):
        ax.text(targ + 1, idx - 0.16, f"{targ}", va="center", fontsize=11.5, color="#2D2D2D")
        ax.text(curr + 1, idx + 0.16, f"{curr}", va="center", fontsize=11.5, color="#2D2D2D")
        note_ax.text(0.02, 0.88 - idx * 0.24, note, va="center", fontsize=10.5, color="#333333", wrap=True)

    ax.legend(loc="lower right", frameon=False, fontsize=11)
    plt.savefig(out_path, dpi=DOCX_CHART_DPI, bbox_inches="tight", facecolor="white")
    plt.close()
    return out_path

def add_pillar3_summary_dashboard(doc):
    chart_path = create_pillar3_summary_chart()
    optimized_path = optimize_image_for_display(chart_path, 6.7)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(optimized_path, width=Inches(6.7))
    doc.add_paragraph("")

def resolve_existing_image(*candidates):
    for candidate in candidates:
        if not candidate:
            continue
        path = candidate
        if not os.path.isabs(path):
            path = os.path.join(DATA_DIR, candidate)
            if not os.path.exists(path):
                path = os.path.join(CHARTS_DIR, candidate)
        if os.path.exists(path):
            return path
    return ""

def _cro_slug(area):
    normalized = re.sub(r"[^a-z0-9]+", "_", str(area).lower()).strip("_")
    return normalized or "cro_area"

def _cro_issue_palette(area):
    if area == "Primary CTA":
        return {"accent": (46, 103, 209), "warn": (217, 83, 79), "soft": (234, 242, 255)}
    if area == "Above-the-Fold Value Proposition":
        return {"accent": (47, 166, 200), "warn": (240, 135, 36), "soft": (236, 248, 251)}
    if area == "Trust Signals":
        return {"accent": (46, 156, 98), "warn": (217, 83, 79), "soft": (235, 248, 241)}
    return {"accent": (124, 92, 214), "warn": (240, 135, 36), "soft": (243, 240, 255)}

def _cro_draw_browser_frame(draw, width, height):
    draw.rounded_rectangle((38, 38, width - 38, height - 38), radius=28, fill=(255, 255, 255), outline=(216, 224, 236), width=2)
    draw.rounded_rectangle((38, 38, width - 38, 100), radius=28, fill=(244, 247, 251))
    draw.rectangle((38, 76, width - 38, 100), fill=(244, 247, 251))
    for idx, color in enumerate([(255, 95, 87), (254, 188, 46), (40, 200, 64)]):
        draw.ellipse((68 + idx * 28, 58, 84 + idx * 28, 74), fill=color)
    draw.rounded_rectangle((170, 54, width - 150, 82), radius=14, fill=(255, 255, 255), outline=(216, 224, 236), width=1)
    draw.text((192, 54), PROSPECT_DOMAIN or "prospect-site.com", font=_load_dashboard_font(18), fill=(102, 112, 133))

def _cro_badge(draw, x, y, text, color):
    font = _load_dashboard_font(18, bold=True)
    text_width = int(draw.textlength(text, font=font))
    draw.rounded_rectangle((x, y, x + text_width + 28, y + 34), radius=14, fill=color)
    draw.text((x + 14, y + 7), text, font=font, fill=(255, 255, 255))

def _cro_wrapped_lines(draw, text, font, max_width):
    lines = []
    current = ""
    for word in str(text).replace("\n", " ").split():
        candidate = (current + " " + word).strip()
        if not current or draw.textlength(candidate, font=font) <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [""]

def _cro_multiline(draw, text, box, size, bold=False, fill=(31, 42, 68), spacing=8, max_lines=None):
    font = _load_dashboard_font(size, bold=bold)
    lines = _cro_wrapped_lines(draw, text, font, box[2] - box[0])
    if max_lines and len(lines) > max_lines:
        lines = lines[:max_lines]
        lines[-1] = lines[-1].rstrip(" ,.;:") + "..."
    draw.multiline_text((box[0], box[1]), "\n".join(lines), font=font, fill=fill, spacing=spacing)
    return font, lines

def _rgb_tuple_from_hex(hex_color):
    clean = str(hex_color).replace("#", "").strip()
    if len(clean) != 6:
        clean = "2E5090"
    return tuple(int(clean[i:i + 2], 16) for i in (0, 2, 4))

def _company_slug():
    raw = business_data.get("company_name") or PROSPECT_NAME or PROSPECT_DOMAIN
    slug = re.sub(r"[^a-z0-9]+", "_", str(raw).lower()).strip("_")
    return slug or "prospect"

def _pick_company_palette():
    key = f"{business_data.get('domain', PROSPECT_DOMAIN)}|{business_data.get('industry', '')}|{PROSPECT_NAME}"
    seed = int(hashlib.md5(key.encode("utf-8")).hexdigest(), 16)
    palettes = [
        {"accent": (46, 107, 230), "secondary": (14, 165, 164), "soft": (235, 244, 255), "muted": (107, 119, 140)},
        {"accent": (11, 104, 155), "secondary": (46, 125, 50), "soft": (235, 246, 250), "muted": (101, 114, 133)},
        {"accent": (34, 94, 168), "secondary": (240, 135, 36), "soft": (240, 245, 252), "muted": (99, 110, 129)},
        {"accent": (71, 85, 182), "secondary": (22, 163, 74), "soft": (240, 242, 255), "muted": (111, 121, 140)},
    ]
    palette = palettes[seed % len(palettes)].copy()
    palette["text"] = (31, 42, 68)
    palette["panel"] = (255, 255, 255)
    palette["border"] = (216, 224, 236)
    palette["track"] = (232, 238, 245)
    return palette

def _find_homepage_screenshot():
    candidates = [
        os.environ.get("HOMEPAGE_SCREENSHOT_PATH", ""),
        os.path.join(DATA_DIR, "homepage_screenshot.png"),
        os.path.join(DATA_DIR, "homepage_screenshot.jpg"),
        os.path.join(CHARTS_DIR, "homepage_screenshot.png"),
    ]
    for candidate in candidates:
        if candidate and os.path.exists(candidate):
            return candidate
    return ""

def _company_visual_context():
    entity = business_data.get("entity_signals", {}) or {}
    nap = entity.get("nap_data", {}) or {}
    return {
        "name": business_data.get("company_name") or PROSPECT_NAME,
        "domain": business_data.get("domain") or PROSPECT_DOMAIN,
        "tagline": business_data.get("tagline") or business_data.get("description", ""),
        "description": business_data.get("description") or "",
        "industry": business_data.get("industry") or "",
        "audience": business_data.get("target_audience") or "",
        "geo": business_data.get("geographic_focus") or "",
        "services": business_data.get("primary_services") or [],
        "usp": business_data.get("unique_selling_points") or [],
        "cta": business_data.get("cta_text") or "Request a Quote",
        "has_testimonials": bool(business_data.get("has_testimonials")),
        "has_contact_form": bool(business_data.get("has_contact_form")),
        "phone": nap.get("phone") or "Call now",
        "palette": _pick_company_palette(),
        "homepage_screenshot": _find_homepage_screenshot(),
        "slug": _company_slug(),
    }

def _fit_cover(image, target_size):
    target_w, target_h = target_size
    ratio = max(target_w / max(image.width, 1), target_h / max(image.height, 1))
    resized = image.resize(
        (max(1, int(round(image.width * ratio))), max(1, int(round(image.height * ratio)))),
        Image.Resampling.LANCZOS,
    )
    left = max(0, (resized.width - target_w) // 2)
    top = max(0, (resized.height - target_h) // 2)
    return resized.crop((left, top, left + target_w, top + target_h))

def _crop_homepage_focus(area, screenshot_path):
    crop_map = {
        "Primary CTA": (0.05, 0.05, 0.86, 0.42),
        "Above-the-Fold Value Proposition": (0.00, 0.00, 1.00, 0.40),
        "Trust Signals": (0.04, 0.34, 0.96, 0.72),
        "Form/Lead Capture": (0.44, 0.18, 0.98, 0.90),
    }
    if not screenshot_path or not os.path.exists(screenshot_path):
        return None
    with Image.open(screenshot_path) as src:
        img = src.convert("RGB")
        left_r, top_r, right_r, bottom_r = crop_map.get(area, (0.05, 0.05, 0.95, 0.85))
        box = (
            int(img.width * left_r),
            int(img.height * top_r),
            int(img.width * right_r),
            int(img.height * bottom_r),
        )
        return _fit_cover(img.crop(box), (968, 560))

def _issue_bullets(text):
    tokens = re.split(r"[.;]|, and |, but | and ", text or "")
    cleaned = []
    for token in tokens:
        piece = token.strip(" ,-")
        if len(piece) >= 10:
            cleaned.append(textwrap.shorten(piece, width=42, placeholder="..."))
        if len(cleaned) == 3:
            break
    return cleaned or ["Issue not yet documented"]

def _hero_title(context):
    services = context.get("services") or []
    geo = context.get("geo") or ""
    if services:
        return f"{services[0]} in {geo}".strip() if geo else services[0]
    if context.get("tagline"):
        return context["tagline"]
    industry = context.get("industry") or "service"
    return f"Trusted {industry.lower()} support"

def _hero_subtitle(context):
    tagline = context.get("tagline") or context.get("description") or ""
    if tagline:
        return textwrap.shorten(tagline, width=110, placeholder="...")
    audience = context.get("audience") or "local customers"
    geo = context.get("geo") or "your market"
    return f"Designed to help {audience.lower()} quickly understand the offer and take action across {geo}."

def _draw_before_cro_panel(draw, img, area, context, finding):
    palette = context["palette"]
    screenshot = _crop_homepage_focus(area, context.get("homepage_screenshot"))
    if screenshot is None:
        screenshot = Image.new("RGB", (968, 560), (246, 249, 252))
        screenshot_draw = ImageDraw.Draw(screenshot)
        screenshot_draw.rounded_rectangle((24, 24, 944, 536), radius=28, fill=(255, 255, 255), outline=palette["border"], width=2)
        screenshot_draw.text((58, 60), context["name"], font=_load_dashboard_font(24, bold=True), fill=palette["accent"])
        _cro_multiline(screenshot_draw, _hero_title(context), (58, 116, 884, 220), 36, bold=True, fill=palette["text"], max_lines=2)
        _cro_multiline(screenshot_draw, _hero_subtitle(context), (58, 236, 884, 316), 22, fill=palette["muted"], max_lines=3, spacing=6)
    img.paste(screenshot, (88, 240))
    draw.rounded_rectangle((1048, 240, 1496, 800), radius=24, fill=(255, 255, 255), outline=palette["border"], width=2)
    draw.text((1086, 280), "Observed on current site", font=_load_dashboard_font(28, bold=True), fill=palette["text"])
    for idx, item in enumerate(_issue_bullets(finding.get("current_status", ""))):
        y = 348 + idx * 112
        draw.ellipse((1088, y + 8, 1104, y + 24), fill=(236, 92, 65))
        _cro_multiline(draw, item, (1128, y, 1456, y + 70), 22, fill=palette["muted"], max_lines=3, spacing=6)
    draw.rounded_rectangle((1088, 690, 1456, 760), radius=18, fill=palette["soft"])
    _cro_multiline(draw, textwrap.shorten(context["domain"], width=34, placeholder="..."), (1120, 714, 1428, 744), 20, bold=True, fill=palette["accent"], max_lines=1)

def _draw_after_cro_panel(draw, area, context, finding):
    palette = context["palette"]
    services = context.get("services") or ["Priority service"]
    usps = context.get("usp") or ["Fast response", "Clear pricing", "Experienced team"]
    cta = context.get("cta") or "Request a Quote"
    name = context["name"]

    if area == "Primary CTA":
        _cro_multiline(draw, _hero_title(context), (96, 250, 824, 350), 44, bold=True, fill=palette["text"], max_lines=2)
        _cro_multiline(draw, _hero_subtitle(context), (96, 370, 840, 452), 24, fill=palette["muted"], max_lines=3, spacing=6)
        draw.rounded_rectangle((96, 500, 338, 566), radius=18, fill=palette["accent"])
        draw.text((142, 520), textwrap.shorten(cta, width=18, placeholder="..."), font=_load_dashboard_font(24, bold=True), fill=(255, 255, 255))
        draw.rounded_rectangle((360, 500, 586, 566), radius=18, fill=(255, 255, 255), outline=palette["accent"], width=3)
        draw.text((414, 520), "Call Now", font=_load_dashboard_font(24, bold=True), fill=palette["accent"])
        draw.rounded_rectangle((930, 246, 1480, 728), radius=28, fill=palette["soft"])
        for idx, text in enumerate((services + usps)[:5]):
            y = 318 + idx * 78
            draw.rounded_rectangle((972, y, 1426, y + 50), radius=16, fill=(255, 255, 255))
            _cro_multiline(draw, text, (1000, y + 11, 1392, y + 43), 21, bold=True, fill=palette["text"], max_lines=2, spacing=4)
    elif area == "Above-the-Fold Value Proposition":
        title = context.get("tagline") or _hero_title(context)
        _cro_multiline(draw, title, (96, 250, 880, 336), 42, bold=True, fill=palette["text"], max_lines=2)
        desc = context.get("description") or finding.get("opportunity", "")
        _cro_multiline(draw, desc, (96, 360, 860, 454), 23, fill=palette["muted"], max_lines=3, spacing=6)
        for idx, text in enumerate((usps + services)[:4]):
            x = 96 + (idx % 2) * 342
            y = 506 + (idx // 2) * 104
            draw.rounded_rectangle((x, y, x + 312, y + 76), radius=18, fill=(247, 250, 253), outline=palette["border"], width=2)
            _cro_multiline(draw, text, (x + 20, y + 18, x + 286, y + 56), 21, bold=True, fill=palette["text"], max_lines=2, spacing=4)
        draw.rounded_rectangle((912, 274, 1482, 702), radius=28, fill=(255, 255, 255), outline=palette["border"], width=2)
        draw.text((952, 322), "Why this is clearer", font=_load_dashboard_font(28, bold=True), fill=palette["text"])
        for idx, text in enumerate(["Sharper message hierarchy", "Clearer local relevance", "Immediate proof of value"]):
            yy = 404 + idx * 88
            draw.ellipse((954, yy, 972, yy + 18), fill=palette["secondary"])
            _cro_multiline(draw, text, (1000, yy - 8, 1438, yy + 34), 22, fill=palette["muted"], max_lines=2, spacing=4)
    elif area == "Trust Signals":
        _cro_multiline(draw, f"Why customers choose {name}", (96, 248, 980, 330), 42, bold=True, fill=palette["text"], max_lines=2)
        _cro_multiline(draw, "Trust elements are placed close to the decision point to reduce hesitation and reinforce credibility.", (96, 356, 990, 430), 23, fill=palette["muted"], max_lines=2, spacing=6)
        trust_cards = []
        if context.get("has_testimonials"):
            trust_cards.append("Verified customer reviews")
        trust_cards.extend(usps[:2])
        trust_cards.append(context.get("geo") or "Local coverage")
        for idx, text in enumerate(trust_cards[:4]):
            x = 96 + idx * 342
            draw.rounded_rectangle((x, 472, x + 286, 538), radius=18, fill=(247, 250, 253), outline=palette["border"], width=2)
            _cro_multiline(draw, text, (x + 18, 490, x + 262, 524), 20, bold=True, fill=palette["text"], max_lines=2, spacing=4)
        for idx, quote in enumerate([
            f"Clear communication and fast help from {name}.",
            "Booking felt simple and reassuring.",
            "Professional service with visible proof points.",
        ]):
            x = 96 + idx * 474
            draw.rounded_rectangle((x, 590, x + 414, 778), radius=24, fill=(255, 255, 255), outline=palette["border"], width=2)
            draw.text((x + 24, 618), "5-star feedback", font=_load_dashboard_font(24, bold=True), fill=palette["secondary"])
            _cro_multiline(draw, quote, (x + 24, 658, x + 382, 748), 21, fill=palette["text"], max_lines=4, spacing=6)
    else:
        _cro_multiline(draw, f"Fast quote flow for {name}", (96, 248, 860, 332), 42, bold=True, fill=palette["text"], max_lines=2)
        _cro_multiline(draw, "The form is simplified around only the information needed to start the conversation.", (96, 356, 860, 432), 23, fill=palette["muted"], max_lines=2, spacing=6)
        draw.rounded_rectangle((96, 450, 782, 794), radius=28, fill=(247, 250, 253), outline=palette["border"], width=2)
        for idx, field in enumerate(["Name", "Phone", "Suburb / Location", "How can we help?"]):
            y = 516 + idx * 66
            draw.text((138, y - 26), field, font=_load_dashboard_font(18, bold=True), fill=palette["muted"])
            draw.rounded_rectangle((138, y, 740, y + 44), radius=14, fill=(255, 255, 255), outline=palette["border"], width=2)
        draw.rounded_rectangle((138, 734, 404, 786), radius=18, fill=palette["accent"])
        draw.text((186, 747), textwrap.shorten(cta, width=18, placeholder="..."), font=_load_dashboard_font(24, bold=True), fill=(255, 255, 255))
        draw.rounded_rectangle((912, 470, 1480, 756), radius=26, fill=(255, 255, 255), outline=palette["border"], width=2)
        draw.text((950, 510), "Conversion improvements", font=_load_dashboard_font(28, bold=True), fill=palette["text"])
        for idx, text in enumerate(["Fewer form fields", "Clearer CTA priority", "Mobile-friendly spacing", "More confidence before submit"]):
            yy = 574 + idx * 46
            draw.ellipse((952, yy, 968, yy + 16), fill=palette["secondary"])
            _cro_multiline(draw, text, (994, yy - 8, 1440, yy + 26), 21, fill=palette["muted"], max_lines=2, spacing=4)

def generate_cro_section_mockup(area, mode="before", finding=None, context=None):
    context = context or _company_visual_context()
    finding = finding or {}
    slug = _cro_slug(area)
    out_path = os.path.join(DATA_DIR, f"cro_{context['slug']}_{slug}_{mode}.png")
    img = Image.new("RGB", (1600, 900), (251, 252, 254))
    draw = ImageDraw.Draw(img)
    palette = context["palette"]
    _cro_draw_browser_frame(draw, 1600, 900)

    title_suffix = "Before CRO" if mode == "before" else "After CRO"
    subtitle = (
        f"Focused view of {context['name']}'s current {area.lower()} experience."
        if mode == "before"
        else f"Optimized concept tailored to {context['name']}'s offer and brand cues."
    )
    draw.text((56, 118), f"{context['name']} - {area} - {title_suffix}", font=_load_dashboard_font(34, bold=True), fill=palette["text"])
    draw.text((56, 160), subtitle, font=_load_dashboard_font(22), fill=palette["muted"])
    draw.rounded_rectangle((56, 206, 1544, 828), radius=26, fill=(255, 255, 255), outline=palette["border"], width=2)

    if mode == "before":
        _draw_before_cro_panel(draw, img, area, context, finding)
    else:
        _draw_after_cro_panel(draw, area, context, finding)

    img.save(out_path, dpi=DOCX_IMAGE_DPI)
    return out_path


def add_inline_image(doc, path, width=6.35):
    if not path or not os.path.exists(path):
        return None
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(optimize_image_for_display(path, width), width=Inches(width))
    doc.add_paragraph("")
    return p

def add_cro_focus_section(doc, finding):
    context = _company_visual_context()
    area = finding.get("area", "CRO Area")
    current_status = finding.get("current_status", "Not verified.")
    opportunity = finding.get("opportunity", "Awaiting assessment.")

    add_heading_styled(doc, area, 3)
    add_body(doc, f"Current Status: {current_status}")
    add_body(doc, f"Opportunity: {opportunity}")

    before_path = generate_cro_section_mockup(area, mode="before", finding=finding, context=context)
    after_path = generate_cro_section_mockup(area, mode="after", finding=finding, context=context)

    add_body(doc, "Before CRO:", bold=True, color=NAVY)
    add_inline_image(doc, before_path, width=6.35)
    add_body(doc, "After CRO:", bold=True, color=NAVY)
    add_inline_image(doc, after_path, width=6.35)

def justify_all_content(doc):
    for para in doc.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_visual_enrichment_block(doc, title, candidates, width=6.0):
    inserted_any = False
    for filename, caption in candidates:
        path = os.path.join(CHARTS_DIR, filename)
        if os.path.exists(path):
            if not inserted_any:
                add_heading_styled(doc, title, 2)
            add_chart_with_caption(doc, filename, caption, width=width)
            inserted_any = True
    return inserted_any

def build_concise_executive_summary(text, max_sentences=3, max_chars=560):
    if not text:
        return ""
    normalized = " ".join(str(text).split())
    sentences = re.split(r'(?<=[.!?])\s+', normalized)
    concise = []
    total_len = 0
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        next_len = total_len + (1 if concise else 0) + len(sentence)
        if len(concise) >= max_sentences or next_len > max_chars:
            break
        concise.append(sentence)
        total_len = next_len
    if not concise:
        trimmed = normalized[:max_chars].rstrip()
        return trimmed + ("..." if len(trimmed) < len(normalized) else "")
    summary = " ".join(concise).strip()
    if len(summary) < len(normalized) and not summary.endswith((".", "!", "?")):
        summary += "..."
    return summary

def replace_placeholder_tokens(doc, replacements):
    pattern = re.compile("|".join(re.escape(k) for k in replacements.keys()), re.IGNORECASE)
    if not replacements:
        return

    def _replace_text(text):
        return pattern.sub(lambda m: replacements.get(m.group(0), replacements.get(m.group(0).lower(), m.group(0))), text)

    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                run.text = _replace_text(run.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = _replace_text(run.text)

def add_page_break(doc):
    # Insert exactly one page break before each main chapter transition.
    if not doc.paragraphs:
        return
    # Remove trailing empty paragraphs before checking/inserting a break.
    while len(doc.paragraphs) > 1:
        last_para = doc.paragraphs[-1]
        has_text = bool(last_para.text.strip())
        has_drawing = bool(last_para._element.findall(".//w:drawing", namespaces=last_para._element.nsmap))
        has_picture = bool(last_para._element.findall(".//w:pict", namespaces=last_para._element.nsmap))
        has_page_break = bool(
            any(br.get(qn("w:type")) == "page" for br in last_para._element.findall(".//w:br", namespaces=last_para._element.nsmap))
        )
        if has_text or has_drawing or has_picture or has_page_break:
            break
        para = doc.paragraphs[-1]
        para._element.getparent().remove(para._element)
    last_para = doc.paragraphs[-1]
    br_nodes = last_para._element.findall(".//w:br", namespaces=last_para._element.nsmap)
    for br in br_nodes:
        if br.get(qn("w:type")) == "page":
            return
    doc.add_page_break()

def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = "© Copyright 2026 TrafficRadius, All rights reserved.\t\tPage "
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(8)
        run.font.color.rgb = LIGHT_GREY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ══════════════════════════════════════════════════════════════
# DOCUMENT CREATION
# ══════════════════════════════════════════════════════════════

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = DARK_GREY
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(6)
pf.line_spacing = 1.15

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

add_footer(doc.sections[0])

# ── COVER PAGE ───────────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run()
if os.path.exists(LOGO_PATH):
    run.add_picture(optimize_image_for_display(LOGO_PATH, 2.5), width=Inches(2.5))

for _ in range(3):
    doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run(PROSPECT_NAME)
run.font.size = Pt(36)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Integrated SEO, AEO & GEO")
run.font.size = Pt(28)
run.font.color.rgb = BLUE
run.font.name = 'Calibri'
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Prospect Audit & Growth Strategy")
run.font.size = Pt(22)
run.font.color.rgb = BLUE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("A Data-Driven Blueprint for Multi-Channel Search Dominance")
run.font.size = Pt(14)
run.font.color.rgb = MID_GREY
run.font.name = 'Calibri'
run.italic = True

for _ in range(4):
    doc.add_paragraph("")

add_divider(doc)

meta_table = doc.add_table(rows=4, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
meta_data = [
    ("Prepared For:", PROSPECT_NAME),
    ("Prepared By:", "TrafficRadius — trafficradius.com.au"),
    ("Date:", PROSPECT_DATE),
    ("Classification:", "Confidential — Client Use Only"),
]
for i, (label, value) in enumerate(meta_data):
    cell_l = meta_table.rows[i].cells[0]
    cell_l.text = ""
    run_l = cell_l.paragraphs[0].add_run(label)
    run_l.bold = True
    run_l.font.size = Pt(11)
    run_l.font.color.rgb = NAVY
    run_l.font.name = 'Calibri'
    cell_l.width = Inches(1.8)
    cell_r = meta_table.rows[i].cells[1]
    cell_r.text = ""
    run_r = cell_r.paragraphs[0].add_run(value)
    run_r.font.size = Pt(11)
    run_r.font.color.rgb = DARK_GREY
    run_r.font.name = 'Calibri'
apply_full_grid_borders(meta_table, color=BLACK_HEX, sz="10")

doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("© Copyright 2026 TrafficRadius, All rights reserved.")
run.font.size = Pt(9)
run.font.color.rgb = LIGHT_GREY
run.font.name = 'Calibri'

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 1: EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "1. Executive Summary", 1)

executive_summary_text = narrative_data.get(
    "executive_summary",
    f"This document presents a comprehensive, multi-layer digital audit and growth strategy for {PROSPECT_NAME}. "
    "Our analysis spans the three pillars of modern search — Traditional SEO, Answer Engine Optimization (AEO), "
    "and Generative Engine Optimization (GEO) — to identify opportunities across every channel where "
    "potential customers are searching for solutions."
)

add_body(doc, build_concise_executive_summary(executive_summary_text))

add_callout(doc,
    "Key Insight: [PROSPECT_NAME] currently captures less than X% of the total addressable "
    "search market in its core verticals, with significant untapped opportunities in answer "
    "engines and AI-powered search platforms."
)

add_body(doc,
    "Our strategy is built on an integrated framework that addresses all three search layers "
    "through a phased, 90-day action plan. This approach ensures sustainable, long-term growth "
    "across traditional search, featured snippets, and AI-generated answers."
)

# Add a single executive dashboard at the bottom of the same page.
add_chart_with_caption(
    doc,
    "integrated_scorecard.png",
    "Dashboard: Current readiness across SEO, AEO, and GEO.",
    width=5.8
)
add_executive_snapshot_dashboard(
    doc,
    "Executive Snapshot Dashboard",
    [
        ("Organic Keywords", market_data.get("prospect", {}).get("organic_keywords", "N/A")),
        ("Monthly Traffic", market_data.get("prospect", {}).get("organic_traffic", "N/A")),
        ("Traffic Value", market_data.get("prospect", {}).get("organic_traffic_value", "N/A")),
        ("Overall Score", audit_data.get("scorecard", {}).get("overall_score", "N/A")),
    ]
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 2: THE THREE LAYERS OF MODERN SEARCH
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "2. The Three Layers of Modern Search", 1)

add_body(doc,
    "The search landscape has fundamentally changed. Today, businesses must be visible not just "
    "in traditional Google results, but across three distinct layers of search. Understanding "
    "these layers is critical to building an effective digital strategy."
)

add_heading_styled(doc, "Layer 1: SEO — Traditional Search Engine Optimization", 2)
add_body(doc,
    "Traditional SEO focuses on ranking in organic search results on Google, Bing, and other "
    "search engines. This remains the foundation of digital visibility and drives the majority "
    "of high-intent commercial traffic."
)

add_heading_styled(doc, "Layer 2: AEO — Answer Engine Optimization", 2)
add_body(doc,
    "AEO focuses on winning Featured Snippets, People Also Ask (PAA) boxes, and voice search "
    "results. These positions capture significant click-through rates and establish the brand "
    "as the definitive answer to common industry questions."
)

add_heading_styled(doc, "Layer 3: GEO — Generative Engine Optimization", 2)
add_body(doc,
    "GEO is the newest frontier, focusing on being cited and referenced by AI-powered search "
    "platforms such as ChatGPT, Google AI Overviews, Perplexity, and Claude. As more users "
    "turn to AI for answers, GEO ensures the brand is part of the conversation."
)

#add_heading_styled(doc, "Three Layers of Modern Search Dashboard", 2)
add_chart_with_caption(
    doc,
    "three_layer_overview.png",
    "Dashboard: Visibility and readiness across SEO, AEO, and GEO layers.",
    width=5.8
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 3: BUSINESS & MARKET CONTEXT
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "3. Business & Market Context", 1)

add_heading_styled(doc, "Company Profile", 2)
add_body(doc,
    narrative_data.get("company_profile", 
        f"TrafficRadius identified {PROSPECT_NAME} as a key operator within its niche, requiring an enhanced discovery posture across SEO, AEO, and GEO platforms."
    )
)

add_heading_styled(doc, "Digital Maturity Assessment", 2)
add_body(doc,
    build_concise_executive_summary(
        narrative_data.get(
            "digital_maturity_assessment",
            "Currently, the prospect demonstrates a foundational baseline of SEO assets, however, advanced schema, comprehensive backlink integration, and strict technical AI optimization parameters are lacking. Implementation of the outlined strategies will exponentially grow market share."
        ),
        max_sentences=2,
        max_chars=340,
    )
)

# Integrated market visibility view within the chapter flow.
add_body(
    doc,
    "The following market visibility dashboard highlights where current search demand is concentrated across the three core discovery layers, helping frame the scale of opportunity before strategic execution begins."
)
create_search_demand_dashboard()
add_chart_with_caption(
    doc,
    "search_demand_by_cluster.png",
    "Dashboard: Total monthly search demand by category.",
    width=5.8
)
add_body(
    doc,
    f"Overall, {PROSPECT_NAME} is operating in a market where a small set of high-value clusters drives the greatest upside. This makes phased prioritization essential, starting with the strongest commercial categories and then extending into answer-engine and AI-led visibility."
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 4: CURRENT DIGITAL PERFORMANCE
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "4. Current Digital Performance", 1)

add_body(doc,
    f"The following metrics provide a data-driven baseline of {PROSPECT_NAME}'s current "
    "organic search performance, sourced from SEMrush."
)

prospect_data = market_data.get("prospect", {})
add_branded_table(doc,
    headers=["Metric", "Current Value"],
    rows=[
        ["Domain Authority Score", str(prospect_data.get("authority_score", "N/A"))],
        ["Organic Keywords", str(prospect_data.get("organic_keywords", "N/A"))],
        ["Monthly Organic Traffic", str(prospect_data.get("organic_traffic", "N/A"))],
        ["Organic Traffic Value", f"${prospect_data.get('organic_traffic_value', '0')}"],
        ["Backlinks", f"{prospect_data.get('backlinks', 'N/A'):,}" if isinstance(prospect_data.get('backlinks'), int) else prospect_data.get('backlinks', 'N/A')],
        ["Referring Domains", f"{prospect_data.get('referring_domains', 'N/A'):,}" if isinstance(prospect_data.get('referring_domains'), int) else prospect_data.get('referring_domains', 'N/A')],
    ]
)

add_heading_styled(doc, "Performance Insights Snapshot", 2)
add_body(
    doc,
    "The current performance baseline indicates clear upside potential when visibility and conversion "
    "opportunities are prioritized in high-value keyword clusters."
)
for point in [
    "Authority and referring domain growth can strengthen competitiveness in core service categories.",
    "Organic traffic value indicates room for faster ROI through targeted optimization of top-intent pages.",
    "A focused execution plan can convert existing visibility into measurable lead and revenue gains.",
]:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(point)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK_GREY

create_traffic_value_opportunity_dashboard()
add_chart_with_caption(
    doc,
    "traffic_value_opportunity.png",
    "Dashboard: Estimated monthly traffic value opportunity by category.",
    width=5.8
)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 5: COMPETITIVE LANDSCAPE
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "5. Competitive Landscape", 1)

add_body(doc,
    narrative_data.get("competitive_landscape_analysis", 
        f"Understanding the competitive landscape is critical to identifying where {PROSPECT_NAME} "
        "can gain the most ground across all three search layers."
    )
)

comp_rows = [
    [
        PROSPECT_NAME,
        str(market_data.get("prospect", {}).get("organic_keywords", "N/A")),
        str(market_data.get("prospect", {}).get("organic_traffic", "N/A")),
        str(market_data.get("prospect", {}).get("organic_traffic_value", "N/A")),
        f"{market_data.get('prospect', {}).get('backlinks', 'N/A'):,}" if isinstance(market_data.get('prospect', {}).get('backlinks'), int) else str(market_data.get('prospect', {}).get('backlinks', 'N/A')),
    ]
]
for comp in market_data.get("competitors", [])[:3]:
    backlinks = comp.get("backlinks", "N/A")
    comp_rows.append([
        comp.get("domain", ""),
        str(comp.get("organic_keywords", "N/A")),
        str(comp.get("organic_traffic", "N/A")),
        str(comp.get("organic_traffic_value", "N/A")),
        f"{backlinks:,}" if isinstance(backlinks, int) else str(backlinks),
    ])

if len(comp_rows) == 1:
    comp_rows.extend([
        ["Competitor 1", "N/A", "N/A", "N/A", "N/A"],
        ["Competitor 2", "N/A", "N/A", "N/A", "N/A"],
        ["Competitor 3", "N/A", "N/A", "N/A", "N/A"],
    ])

add_branded_table(doc,
    headers=["Competitor", "Organic Keywords", "Monthly Traffic", "Traffic Value", "Backlinks"],
    rows=comp_rows
)

create_competitive_landscape_dashboard()
add_chart_image(doc, "competitive_landscape.png")

if shadow_data or market_data.get("competitors"):
    add_heading_styled(doc, "5.1 Value Proposition Gap Analysis", 2)
    add_body(doc, shadow_data.get("overall_gap_summary", "Competitor analysis indicates several missed opportunities."))
    for idx, (competitor, competitor_data, gap) in enumerate(_resolve_gap_entries(limit=3), start=1):
        insight = (
            f"{competitor} insight: {textwrap.shorten(gap.get('what_they_have', '').strip(), width=110, placeholder='...')} "
            f"The response priority is to {textwrap.shorten(gap.get('how_to_beat_it', '').strip(), width=120, placeholder='...')}"
        )
        add_body(doc, insight)
        gap_dashboard_path = create_value_gap_competitor_dashboard(competitor, competitor_data, gap, idx)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(optimize_image_for_display(gap_dashboard_path, 6.5), width=Inches(6.5))
        doc.add_paragraph("")

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 6: INTEGRATED AUDIT — SEO, AEO & GEO
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "6. Integrated Audit — SEO, AEO & GEO", 1)

add_body(doc,
    "Our integrated audit assessed the website across all three search layers. "
    "The scorecard below provides a snapshot of current readiness for each layer."
)

create_integrated_scorecard_dashboard()
add_chart_image(doc, "integrated_scorecard.png")

add_heading_styled(doc, "6.1 SEO Audit Findings", 2)
add_body(doc,
    "The following high-impact technical SEO issues were identified during our audit."
)
qw_rows = []
for qw in narrative_data.get("seo_quick_wins", []):
    qw_rows.append(["High", "Technical SEO", qw.get("title", ""), qw.get("description", "")])

if not qw_rows:
    qw_rows = [
        ["Critical", "Technical SEO", "Awaiting Scan", "Initiate technical sweep"],
        ["High", "Technical SEO", "Awaiting Scan", "Initiate technical sweep"],
        ["Medium", "Technical SEO", "Awaiting Scan", "Initiate technical sweep"],
    ]

add_branded_table(doc,
    headers=["Severity", "Area", "Finding", "Recommendation"],
    rows=qw_rows
)

add_heading_styled(doc, "6.2 AEO Audit Findings", 2)
add_callout(doc,
    "AEO focuses on how well the website is structured to win Featured Snippets, "
    "People Also Ask boxes, and voice search results.",
    AEO_HEX
)
aeo_rows = []
for finding in audit_data.get("aeo_findings", []):
    aeo_rows.append([finding.get("area", "AEO"), finding.get("current_status", "N/A"), finding.get("title", ""), finding.get("recommendation", "")])

if not aeo_rows:
    aeo_rows = [["AEO Readiness", "N/A", "No AEO findings tracked.", "Expand schema deployment."]]

add_branded_table(doc,
    headers=["Area", "Status", "Finding", "Recommendation"],
    rows=aeo_rows
)

add_heading_styled(doc, "6.3 GEO Audit Findings", 2)
add_callout(doc,
    "GEO assesses how visible the brand is to AI-powered search engines like "
    "ChatGPT, Perplexity, and Google AI Overviews.",
    GEO_HEX
)
geo_rows = []
for finding in audit_data.get("geo_findings", []):
    geo_rows.append([finding.get("area", "GEO"), finding.get("current_status", "N/A"), finding.get("title", ""), finding.get("recommendation", "")])

if not geo_rows:
    geo_rows = [["GEO Targets", "N/A", "No GEO findings tracked.", "Build Entity Signals."]]

add_branded_table(doc,
    headers=["Area", "Status", "Finding", "Recommendation"],
    rows=geo_rows
)
add_integrated_layer_visual_dashboard(doc)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 7: CRO ASSESSMENT
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "7. CRO Assessment", 1)

add_body(doc,
    "Beyond attracting traffic across all search layers, it is essential that the website "
    "effectively converts visitors into leads or customers."
)

default_cro_findings = [
    {
        "area": "Primary CTA",
        "current_status": "The page lacks any visible CTA. Users are left without guidance on what action to take next, leading to high bounce rates.",
        "opportunity": "Implement a bold, contrasting CTA button with clear, action-oriented text like 'Get Started Now' or 'Learn More' prominently above the fold.",
    },
    {
        "area": "Above-the-Fold Value Proposition",
        "current_status": "There is no value proposition visible. The page does not clearly communicate what makes the offer compelling or distinct.",
        "opportunity": "Craft a compelling headline and subheadline that clearly communicate the unique value and benefits of the offering, placed prominently above the fold.",
    },
    {
        "area": "Trust Signals",
        "current_status": "No trust signals are present. The absence of reviews, badges, or contact information erodes user confidence.",
        "opportunity": "Add visible trust badges, customer testimonials, and a contact number to establish credibility and reassure visitors.",
    },
    {
        "area": "Form/Lead Capture",
        "current_status": "There is no clear lead capture mechanism visible, resulting in missed opportunities for engagement and conversion.",
        "opportunity": "Design a simple, user-friendly lead capture form with minimal fields, strategically placed to capture visitor information without causing friction.",
    },
]

raw_cro_findings = cro_data.get("findings", [])
cro_findings = []
for default in default_cro_findings:
    default_key = re.sub(r"[^a-z0-9]+", "", default["area"].lower())
    matched = None
    for finding in raw_cro_findings:
        finding_key = re.sub(r"[^a-z0-9]+", "", str(finding.get("area", "")).lower())
        if default_key == finding_key or default_key in finding_key or finding_key in default_key:
            matched = finding
            break
    cro_findings.append({
        "area": default["area"],
        "current_status": (matched or {}).get("current_status", default["current_status"]),
        "opportunity": (matched or {}).get("opportunity", default["opportunity"]),
    })

for finding in cro_findings:
    add_cro_focus_section(doc, finding)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 8: KEYWORD OPPORTUNITY ANALYSIS
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "8. Keyword Opportunity Analysis", 1)

add_body(doc,
    "Our multi-layer keyword research identified significant untapped search demand across "
    f"{PROSPECT_NAME}'s core service categories. Keywords are categorized by their primary "
    "search layer — SEO, AEO, or GEO — to inform a targeted strategy."
)

add_heading_styled(doc, "Opportunity by Search Layer", 2)
create_layer_distribution_dashboard()
add_chart_image(doc, "layer_distribution.png")

add_heading_styled(doc, "Search Demand by Category", 2)
add_chart_image(doc, "search_demand_by_cluster.png")

add_heading_styled(doc, "Top Keyword Opportunities", 2)
kw_rows = []
# Grab Top 3 SEO Keywords from semrush payload
for kw in market_data.get("prospect", {}).get("top_keywords", [])[:3]:
    kw_rows.append([kw.get("Keyword", ""), "SEO", "Service", kw.get("Search Volume", "0"), f"${kw.get('CPC', '0')}", kw.get("Competition", "0"), "High"])

# Grab Top 2 AEO Keywords
for kw in market_data.get("aeo_indicators", {}).get("top_question_keywords", [])[:2]:
    kw_rows.append([kw.get("Keyword", ""), "AEO", "FAQ", kw.get("Search Volume", "0"), f"${kw.get('CPC', '0')}", kw.get("Competition", "0"), "Medium"])

# Grab Top 2 GEO Keywords
for kw in market_data.get("geo_indicators", {}).get("top_informational_keywords", [])[:2]:
    kw_rows.append([kw.get("Keyword", ""), "GEO", "Informational", kw.get("Search Volume", "0"), f"${kw.get('CPC', '0')}", kw.get("Competition", "0"), "High"])

if not kw_rows:
    kw_rows = [["Primary Services SEO", "SEO", "General", "10K", "$10", "0.5", "High"]]

add_branded_table(doc,
    headers=["Keyword", "Layer", "Cluster", "Volume", "CPC", "Competition", "Score"],
    rows=kw_rows
)

add_heading_styled(doc, "Opportunity Matrix", 2)
opportunity_points = [
    "High-volume, medium-competition keywords should be prioritized for faster visibility gains.",
    "High-CPC clusters represent immediate commercial upside when supported by conversion-ready pages.",
    "Question-led topics should be mapped to AEO/GEO formats to expand non-traditional search coverage.",
]
for point in opportunity_points:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(point)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK_GREY
create_opportunity_matrix_dashboard()
add_chart_image(doc, "opportunity_matrix.png")
add_opportunity_mix_dashboard(doc)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 9: INTEGRATED STRATEGY — SEO, AEO & GEO
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "9. Integrated Strategy — SEO, AEO & GEO", 1)

add_body(doc,
    "Our recommended strategy addresses all three layers of modern search through "
    "an integrated framework. Each pillar reinforces the others, creating a compounding "
    "effect that maximizes visibility across every search channel."
)

def render_strategy_pillar(doc, title, pillar_data):
    if not pillar_data:
        add_heading_styled(doc, title, 2)
        add_body(doc, "Data unvailable.")
        return
        
    add_heading_styled(doc, title, 2)
    add_body(doc, pillar_data.get("overview", ""))

    if title.startswith("Pillar 1: Technical Foundation & AI Readiness"):
        add_pillar1_summary_dashboard(doc)
    elif title.startswith("Pillar 2: Content & Answer Optimization"):
        add_pillar2_summary_dashboard(doc)
    elif title.startswith("Pillar 3: Authority & Entity Building"):
        add_pillar3_summary_dashboard(doc)
    
    add_heading_styled(doc, "Key Tactical Initiatives", 3)
    for initiative in pillar_data.get("key_initiatives", []):
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(initiative)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK_GREY

    impact_rows = []
    for item in pillar_data.get("impact_matrix", []):
        impact_rows.append([item.get("opportunity", ""), item.get("expected_outcome", "")])
    
    if impact_rows:
        doc.add_paragraph("")
        add_branded_table(doc, headers=["Opportunity Area", "Expected Business Outcome"], rows=impact_rows)

render_strategy_pillar(doc, "Pillar 1: Technical Foundation & AI Readiness", narrative_data.get("integrated_strategy_technical", {}))
render_strategy_pillar(doc, "Pillar 2: Content & Answer Optimization", narrative_data.get("integrated_strategy_content", {}))
render_strategy_pillar(doc, "Pillar 3: Authority & Entity Building", narrative_data.get("integrated_strategy_authority", {}))

'''add_visual_enrichment_block(
    doc,
    "Integrated Opportunity Dashboards",
    [
        ("layer_distribution.png", "Dashboard: Opportunity split across SEO, AEO, and GEO."),
        ("opportunity_matrix.png", "Dashboard: Prioritization matrix for opportunity targeting."),
    ],
    width=5.8
)'''

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 10: CONTENT STRATEGY
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "10. Content Strategy", 1)

add_body(doc,
    "A well-executed content strategy is the engine that drives growth across all three "
    "search layers. Based on our AI's analysis of your search landscape and competitive gaps, we have formulated customized editorial pillars:"
)

content_roadmap = narrative_data.get("content_strategy_roadmap", [])
if content_roadmap:
    overview_dashboard = create_content_strategy_summary_dashboard()
    if overview_dashboard:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(optimize_image_for_display(overview_dashboard, 6.7), width=Inches(6.7))
        add_body(
            doc,
            "This overview separates the content plan into three priority service lanes so demand, schema readiness, and editorial depth can be reviewed together.",
            color=MID_GREY
        )
    for idx, item in enumerate(content_roadmap, start=1):
        add_heading_styled(doc, item.get("topic", "Content Pillar"), 2)
        add_body(doc, item.get("rationale", ""))
        detail_dashboard = create_content_pillar_detail_dashboard(item, idx)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(optimize_image_for_display(detail_dashboard, 6.7), width=Inches(6.7))
        if idx == 1:
            add_body(
                doc,
                "Editorial recommendation: lead with urgent service language, fast response proof, and clear local relevance to match the underlying search intent.",
                color=MID_GREY
            )
        elif idx == 2:
            add_body(
                doc,
                "Conversion angle: combine diagnostic guidance with service-booking prompts so informational searchers can move quickly into enquiry behaviour.",
                color=MID_GREY
            )
        elif idx == 3:
            add_body(
                doc,
                "Content should reassure users with repair expertise, urgency cues, and clear next-step options for both immediate fixes and broader replacement decisions.",
                color=MID_GREY
            )

        add_heading_styled(doc, "Recommended Sub-Topics (H2 Structure)", 3)
        for sub in item.get("sub_topics", []):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(sub)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_GREY
else:
    add_body(doc, "Expand technical and location-based clusters targeting high-intent long-tail keywords.")

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 11: 90-DAY ACTION PLAN
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "11. Action Plan", 1)

add_body(doc,
    "The following phased roadmap outlines the key initiatives for the first 90 days, "
    "with actions organized across all three search layers."
)

action_plan_rows = [
    ["Phase 1\nFoundation",
     "Fix critical technical issues, optimize title tags and meta descriptions",
     "Implement FAQPage schema, add FAQ sections to top service pages",
     "Ensure AI bots are not blocked, implement Organization schema, create llms.txt"],
    ["Phase 2\nContent",
     "Publish cornerstone content for top clusters, optimize existing pages",
     "Create how-to guides for top question keywords, add Q&A structured content",
     "Publish thought leadership content, strengthen author bios and entity signals"],
    ["Phase 3\nAuthority",
     "Launch outreach campaign, amplify top content",
     "Optimize for PAA opportunities, expand FAQ coverage",
     "Build entity consistency across web, pursue digital PR for citations"],
]

add_branded_table(doc,
    headers=["Phase", "SEO Actions", "AEO Actions", "GEO Actions"],
    rows=action_plan_rows
)
add_action_plan_roadmap(doc, [
    {"phase": row[0].replace("\n", " "), "seo": row[1], "aeo": row[2], "geo": row[3]}
    for row in action_plan_rows
])

add_page_break(doc)

# ══════════════════════════════════════════════════════════════
# SECTION 12: WHY TRAFFICRADIUS & NEXT STEPS
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, "12. Why TrafficRadius & Next Steps", 1)

add_heading_styled(doc, "Why TrafficRadius", 2)
add_body(doc,
    "TrafficRadius is a data-driven digital growth agency that specializes in helping "
    "businesses dominate across all three layers of modern search. Our integrated SEO, AEO, "
    "and GEO approach ensures your brand is visible wherever your customers are searching — "
    "whether in traditional results, answer boxes, or AI-generated responses."
)

add_heading_styled(doc, "Next Steps", 2)
add_body(doc,
    "We are excited about the opportunity to partner with " + PROSPECT_NAME + " and help "
    "you achieve multi-channel search dominance. To get started:"
)

add_branded_table(doc,
    headers=["Step", "Action", "Timeline"],
    rows=[
        ["1", "Review this strategy document and the accompanying data workbook", "This week"],
        ["2", "Schedule a strategy call to discuss findings and answer questions", "Next week"],
        ["3", "Approve the engagement and kick off the project", "Week 3"],
    ]
)
add_trafficradius_next_steps_extension(doc)

# ── Save ─────────────────────────────────────────────────────

# Ensure dynamic prospect naming replaces legacy placeholders everywhere in document text.
replace_placeholder_tokens(doc, {
    "[PROSPECT_NAME]": PROSPECT_NAME,
    "{PROSPECT_NAME}": PROSPECT_NAME,
    "prospect_name": PROSPECT_NAME,
    "PROSPECT_NAME": PROSPECT_NAME,
})

# Apply final alignment rule across all textual content.
justify_all_content(doc)
# Ensure all tables have visible full-grid borders in Word.
for t in doc.tables:
    apply_full_grid_borders(t, color=BLACK_HEX, sz="10")

os.makedirs(OUTPUT_DIR, exist_ok=True)
output_path = os.path.join(OUTPUT_DIR, "Strategy_Document.docx")
doc.save(output_path)
print(f"DOCX saved to: {output_path}")
